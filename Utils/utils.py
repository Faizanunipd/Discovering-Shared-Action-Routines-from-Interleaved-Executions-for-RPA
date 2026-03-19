import os
import random
import numpy as np
import pandas as pd
import scipy.stats as stats
from tabulate import tabulate
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from collections import defaultdict
from IPython.display import display, HTML
from matplotlib.ticker import MaxNLocator
import re

import pm4py
from pm4py.algo.conformance.tokenreplay import algorithm as token_replay
from pm4py.objects.petri_net.importer.variants import pnml as pnml_importer


######################################################################################################################
####                                                   Main Functions                                             ####
######################################################################################################################

# --- Main process function ---
def process_random_logs(logs, document, target_datetime_str="2024-06-07 21:34:25.493914+00:00", variance_criteria='max', specified_routines=None):
    """
    Process random logs with new methodology:
    1. Select three routine types randomly
    2. Shift all selected routine types to target_datetime_str without disturbing inter-trace gaps
    3. Interleave traces with maximum one switch per trace (now across three routines)
    """
    while True:
        # Step 1: Select three routines
        if specified_routines is not None:
            if len(specified_routines) != 3:
                raise ValueError("specified_routines must contain exactly 3 routine types")
            selected_routines = sorted(specified_routines)
        else:
            selected_routines = select_random_routines(logs)
            if len(selected_routines) != 3:
                print(f"Warning: Expected 3 routines but got {len(selected_routines)}")
                if specified_routines is not None:
                    return None, None, None, document, None
                continue
        
        # Step 2: Capture original trace counts BEFORE any processing
        original_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
        traces_per_routine_original = {}
        for i, log in enumerate(original_logs):
            traces_per_routine_original[selected_routines[i]] = log['case:concept:name'].nunique()
        
        # Step 3: Select and renumber cases
        selected_logs = select_and_renumber_cases_all(logs, selected_routines)
        
        # Step 4: Shift all routine types to target_datetime_str without disturbing inter-trace gaps
        shifted_logs = shift_routines_to_target_preserving_gaps(selected_logs, target_datetime_str)
        
        # Step 5: Interleave traces with maximum one switch
        if len(shifted_logs) != 3:
            print(f"Error: Expected 3 logs but got {len(shifted_logs)}")
            if specified_routines is not None:
                return None, None, None, document, None
            continue
        
        final_log = interleave_traces_max_one_switch(shifted_logs[0], shifted_logs[1], shifted_logs[2])
        
        # Trace counts
        num_traces_before = final_log['case:concept:name'].nunique()
        traces_per_routine_before = final_log.groupby('routine_type')['case:concept:name'].nunique()
        num_traces_after = num_traces_before
        traces_per_routine_after = traces_per_routine_before
        
        # Compute interleaving
        interleave_counts = count_trace_interleaving_cases(final_log)

        # Check acceptance conditions
        if (
            num_traces_after >= max(1, int(0.05 * num_traces_before))
        ):
            document = display_and_store_selected_routines(selected_routines, logs, document)
            
            # Collect metadata for iteration summary
            iteration_metadata = {
                'selected_routines': selected_routines,
                'traces_before_renumbering': traces_per_routine_original,
                'traces_after_renumbering': traces_per_routine_after.to_dict(),
                'interleaving_counts': interleave_counts
            }
            
            # Sufficient traces, proceed to output
            # Ensure all required columns exist
            required_columns = ['concept:name', 'time:timestamp', 'case:concept:name', 'routine_type', 'log_number']
            missing_columns = [col for col in required_columns if col not in final_log.columns]
            if missing_columns:
                print(f"Warning: Missing columns in final_log: {missing_columns}")
                print(f"Available columns: {list(final_log.columns)}")
                # Use only available columns
                available_columns = [col for col in required_columns if col in final_log.columns]
                final_log = final_log[available_columns]
            else:
                final_log = final_log[required_columns]
            
            segment_log = final_log.copy()
            unsegment_log = final_log.copy()
            unsegment_log['case:concept:name'] = 1
            
            return segment_log, unsegment_log, selected_routines, document, iteration_metadata
        else:
            if specified_routines is not None:
                return None, None, None, document, None
            continue


######################################################################################################################
####                                            Functions Called by process_random_logs                           ####
######################################################################################################################

def select_random_routines(logs):
    """Select random routines based on trace count and activity criteria"""
    # Static variable to track if this is the first call
    if not hasattr(select_random_routines, 'first_call'):
        select_random_routines.first_call = True
    
    # Calculate trace counts and unique activities for each routine type
    trace_counts = logs.groupby('routine_type')['case:concept:name'].nunique()
    unique_activities = logs.groupby('routine_type')['concept:name'].nunique()
    
    # Only print detailed info on first call
    if select_random_routines.first_call:
        print("*"*50)
        all_trace_couts = sum(logs.groupby('routine_type')['case:concept:name'].nunique())
        print(f"Total routine types before filtering {len(trace_counts)} with trace count {all_trace_couts}")
        
        # Filter by trace count
        trace_mean = int(sum(trace_counts) / len(trace_counts))
        # trace_median = int(trace_counts.median())
        # trace_first_quartile = int(np.percentile(trace_counts, 75))
        trace_eligible = trace_counts[trace_counts > trace_mean].index.tolist()
        print(f"Trace mean: {trace_mean:.1f}")
        print(f"Routine types after trace count filter (> {trace_mean:.1f}): {len(trace_eligible)}")
        
        # Filter by unique activities
        activity_median = unique_activities.median()
        activity_eligible = unique_activities[unique_activities > activity_median].index.tolist()
        print(f"Routine types after activity count filter (> {activity_median:.1f}): {len(activity_eligible)}")
        
        # Get intersection of both filters
        eligible_routines = list(set(trace_eligible) & set(activity_eligible))
        print(eligible_routines)
        eligible_trace_couts = sum(logs[logs['routine_type'].isin(eligible_routines)].groupby('routine_type')['case:concept:name'].nunique())
        print(f"Routine types after both filters are {len(eligible_routines)} with {eligible_trace_couts} traces")
        # print(f"Routine types after both filters: {len(eligible_routines)}")
        
        select_random_routines.first_call = False
        print("*"*50)
    else:
        # For subsequent calls, just calculate the filters without printing
        trace_mean = int(sum(trace_counts) / len(trace_counts))
        trace_eligible = trace_counts[trace_counts > trace_mean].index.tolist()
        activity_median = unique_activities.median()
        activity_eligible = unique_activities[unique_activities > activity_median].index.tolist()
        eligible_routines = list(set(trace_eligible) & set(activity_eligible))
    
    if len(eligible_routines) < 3:
        if select_random_routines.first_call:
            print(f"Warning: Only {len(eligible_routines)} routines meet criteria, selecting all available")
        selected_routines = sorted(eligible_routines)
    else:
        selected_routines = sorted(random.sample(eligible_routines, 3))
    
    print(f"\nSelected routines: {selected_routines}")
    return selected_routines


def select_and_renumber_cases_all(logs, selected_routines):
    """Keep all traces from each routine and renumber case IDs sequentially"""
    selected_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
    
    print(f"Original trace counts per routine:")
    for i, log in enumerate(selected_logs):
        print(f"  Routine {selected_routines[i]}: {log['case:concept:name'].nunique()} traces")
    
    next_case_id = 1
    for i, log in enumerate(selected_logs):
        # Get all unique case IDs for this routine (no limiting)
        unique_cases = log['case:concept:name'].unique()
        
        # Create mapping to new sequential integer IDs
        case_mapping = {old: new for old, new in zip(unique_cases, range(next_case_id, next_case_id + len(unique_cases)))}
        log['case:concept:name'] = log['case:concept:name'].map(case_mapping)
        selected_logs[i] = log
        next_case_id += len(unique_cases)
        
        print(f"  Routine {selected_routines[i]} after renumbering: {len(unique_cases)} traces (case IDs: {min(case_mapping.values())}-{max(case_mapping.values())})")
    
    return selected_logs


def select_and_renumber_cases_minimum(logs, selected_routines):
    """Limit each routine to minimum traces and renumber case IDs sequentially"""
    selected_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
    
    print(f"\n=== BEFORE select_and_renumber_cases_minimum ===")
    original_trace_counts = {}
    for i, log in enumerate(selected_logs):
        trace_count = log['case:concept:name'].nunique()
        original_trace_counts[selected_routines[i]] = trace_count
        print(f"  Routine {selected_routines[i]}: {trace_count} traces")
    
    # Find the minimum number of traces among all logs
    min_traces = min(original_trace_counts.values())
    print(f"Minimum traces across all routines: {min_traces}")
    
    next_case_id = 1
    final_trace_counts = {}
    
    for i, log in enumerate(selected_logs):
        # Select first min_traces case IDs
        unique_cases = log['case:concept:name'].unique()
        selected_case_ids = unique_cases[:min_traces]
        
        log = log[log['case:concept:name'].isin(selected_case_ids)].copy()
        
        # Create mapping to new sequential integer IDs
        case_mapping = {old: new for old, new in zip(selected_case_ids, range(next_case_id, next_case_id + len(selected_case_ids)))}
        log['case:concept:name'] = log['case:concept:name'].map(case_mapping)
        selected_logs[i] = log
        next_case_id += len(selected_case_ids)
        
        # Get the actual final trace count after filtering and renumbering
        final_trace_count = log['case:concept:name'].nunique()
        final_trace_counts[selected_routines[i]] = final_trace_count
        
        print(f"  DEBUG - Routine {selected_routines[i]}: Original={original_trace_counts[selected_routines[i]]}, Selected={len(selected_case_ids)}, Final={final_trace_count}")
    
    print(f"\n=== AFTER select_and_renumber_cases_minimum ===")
    for i, routine in enumerate(selected_routines):
        original_count = original_trace_counts[routine]
        final_count = final_trace_counts[routine]
        reduction = original_count - final_count
        reduction_percent = (reduction / original_count * 100) if original_count > 0 else 0
        print(f"  Routine {routine}: {original_count} → {final_count} traces (reduced by {reduction}, {reduction_percent:.1f}%)")
    
    print(f"\n=== SUMMARY ===")
    total_original = sum(original_trace_counts.values())
    total_final = sum(final_trace_counts.values())
    total_reduction = total_original - total_final
    print(f"  Total traces: {total_original} → {total_final} (reduced by {total_reduction})\n")
    
    return selected_logs


def adjust_timestamps_to_target(selected_logs, target_datetime_str):
    """Adjust timestamps to target datetime"""
    adjusted_logs = []
    target_datetime = pd.to_datetime(target_datetime_str)
    for log in selected_logs:
        timestamps = pd.to_datetime(log['time:timestamp'])
        time_deltas = timestamps - timestamps.min()
        log = log.copy()
        log['time:timestamp'] = target_datetime + time_deltas
        adjusted_logs.append(log)
    return adjusted_logs


def force_first_trace_to_target(adjusted_logs, target_datetime_str):
    """Force first trace to start at target datetime"""
    target_datetime = pd.to_datetime(target_datetime_str)
    new_logs = []
    for log in adjusted_logs:
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min()
        first_event_time = case_starts.min()
        shift = target_datetime - first_event_time
        log = log.copy()
        log['time:timestamp'] = pd.to_datetime(log['time:timestamp']) + shift
        new_logs.append(log)
    return new_logs


def shift_routines_to_target_preserving_gaps(selected_logs, target_datetime_str):
    """
    Shift all routines to target_datetime_str while preserving inter-trace gaps.
    
    Args:
        selected_logs: List of DataFrames, one per routine type
        target_datetime_str: Target datetime string
        
    Returns:
        List of DataFrames with adjusted timestamps preserving gaps
    """
    target_datetime = pd.to_datetime(target_datetime_str)
    shifted_logs = []
    
    for log in selected_logs:
        log = log.copy()
        log['time:timestamp'] = pd.to_datetime(log['time:timestamp'])
        
        # Get all traces (cases) sorted by their start time
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min().sort_values()
        case_ids = case_starts.index.tolist()
        
        if not case_ids:
            shifted_logs.append(log)
            continue
        
        # Calculate original gaps between consecutive traces
        original_gaps = []
        for i in range(len(case_ids) - 1):
            prev_case = log[log['case:concept:name'] == case_ids[i]]
            curr_case = log[log['case:concept:name'] == case_ids[i + 1]]
            prev_end = prev_case['time:timestamp'].max()
            curr_start = curr_case['time:timestamp'].min()
            gap = (curr_start - prev_end).total_seconds()
            original_gaps.append(gap)
        
        # Shift first trace to target datetime
        first_case_id = case_ids[0]
        first_case = log[log['case:concept:name'] == first_case_id]
        first_case_start = first_case['time:timestamp'].min()
        first_shift = target_datetime - first_case_start
        
        # Apply shift to first trace
        mask_first = log['case:concept:name'] == first_case_id
        log.loc[mask_first, 'time:timestamp'] = log.loc[mask_first, 'time:timestamp'] + first_shift
        
        # Shift subsequent traces preserving gaps
        current_end_time = log[log['case:concept:name'] == first_case_id]['time:timestamp'].max()
        
        for i in range(1, len(case_ids)):
            case_id = case_ids[i]
            case_df = log[log['case:concept:name'] == case_id]
            original_start = case_df['time:timestamp'].min()
            
            # Calculate new start time preserving the gap
            gap_seconds = original_gaps[i - 1] if i > 0 else 0
            new_start = current_end_time + pd.Timedelta(seconds=gap_seconds)
            
            # Calculate shift needed
            shift = new_start - original_start
            
            # Apply shift to this trace
            mask = log['case:concept:name'] == case_id
            log.loc[mask, 'time:timestamp'] = log.loc[mask, 'time:timestamp'] + shift
            
            # Update current_end_time for next iteration
            current_end_time = log[log['case:concept:name'] == case_id]['time:timestamp'].max()
        
        shifted_logs.append(log)
    
    return shifted_logs


def interleave_traces_max_one_switch(log1, log2, log3):
    """
    Interleave three logs with maximum one switch per trace.
    
    Process:
    1. Calculate min trace count between all three routines and limit each to that count
    2. Sort traces by start time
    3. For each i-th trace triple:
       - Randomly assign roles: outer, middle, inner (corresponding to Trace 1, Trace 2, Trace 3)
       - Split outer trace at a random point (if possible)
       - Split middle trace at a random point (if possible)
       - Interleave as: outer_before -> middle_before -> inner -> middle_after -> outer_after
         (with small millisecond gaps and timestamp shifts)
    
    The result resembles:
    Trace 1 actions: a1, a2, a3, a4, a5, a6, a7
    Trace 2 actions: b1, b2, b3, b4, b5, b6, b7, b8
    Trace 3 actions: c1, c2, c3, c4, c5, c6
    
    Interleaving:
    a1, a2, a3, b1, b2, b3, b4, b5, c1, c2, c3, c4, c5, c6, b6, b7, b8, a4, a5, a6, a7
    
    Args:
        log1: DataFrame for first routine type
        log2: DataFrame for second routine type
        log3: DataFrame for third routine type
        
    Returns:
        DataFrame with interleaved traces (each trace has max one switch)
    """
    log1 = log1.copy()
    log2 = log2.copy()
    log3 = log3.copy()
    log1['time:timestamp'] = pd.to_datetime(log1['time:timestamp'])
    log2['time:timestamp'] = pd.to_datetime(log2['time:timestamp'])
    log3['time:timestamp'] = pd.to_datetime(log3['time:timestamp'])
    
    # Get all traces from logs sorted by start time
    def get_traces_sorted(log):
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min().sort_values()
        traces = []
        for case_id in case_starts.index:
            trace = log[log['case:concept:name'] == case_id].sort_values('time:timestamp').copy()
            traces.append(trace)
        return traces
    
    traces1 = get_traces_sorted(log1)
    traces2 = get_traces_sorted(log2)
    traces3 = get_traces_sorted(log3)
    
    # If any log has no traces, just return concatenation of non-empty logs (no interleaving possible)
    if not traces1 or not traces2 or not traces3:
        non_empty_logs = [df for df in [log1, log2, log3] if not df.empty]
        if not non_empty_logs:
            return pd.DataFrame()
        result = pd.concat(non_empty_logs, ignore_index=True)
        result = result.sort_values(by='time:timestamp').reset_index(drop=True)
        return result
    
    # Step 1: Calculate minimum trace count and limit all three to that count
    min_trace_count = min(len(traces1), len(traces2), len(traces3))
    print(
        f"Trace counts - Routine1: {len(traces1)}, "
        f"Routine2: {len(traces2)}, Routine3: {len(traces3)}, Min: {min_trace_count}"
    )
    
    traces1 = traces1[:min_trace_count]
    traces2 = traces2[:min_trace_count]
    traces3 = traces3[:min_trace_count]
    
    print(
        f"After limiting - Routine1: {len(traces1)}, "
        f"Routine2: {len(traces2)}, Routine3: {len(traces3)}"
    )
    
    # Store all final traces
    all_final_traces = []
    
    # Calculate a base timestamp for positioning traces
    all_min_times = []
    for trace in traces1 + traces2 + traces3:
        all_min_times.append(trace['time:timestamp'].min())
    base_timestamp = min(all_min_times) if all_min_times else pd.Timestamp.now()
    
    # Track current maximum timestamp to ensure traces don't overlap
    # Initialize to base_timestamp for the first iteration
    current_max_time = base_timestamp
    
    # Use only a few milliseconds gap to avoid overlap
    min_gap_between_iterations = pd.Timedelta(milliseconds=10)
    internal_segment_gap = pd.Timedelta(milliseconds=5)
    
    def place_segment(segment, current_time, add_gap_before=False):
        """Shift a segment so it starts after current_time, optionally adding a small gap."""
        if segment is None or segment.empty:
            return None, current_time
        
        if add_gap_before:
            current_time = current_time + internal_segment_gap
        
        seg_start = segment['time:timestamp'].min()
        shift = current_time - seg_start
        segment = segment.copy()
        segment['time:timestamp'] = segment['time:timestamp'] + shift
        new_current = segment['time:timestamp'].max()
        return segment, new_current
    
    # Step 2: Process traces one by one
    for i in range(min_trace_count):
        trace1 = traces1[i].copy()
        trace2 = traces2[i].copy()
        trace3 = traces3[i].copy()
        
        # For iterations after the first, ensure we start after the previous iteration ends
        # Add a small gap to prevent overlap between consecutive iterations
        if i > 0:
            current_max_time = current_max_time + min_gap_between_iterations
        
        # Randomly assign roles: outer, middle, inner
        indices = [0, 1, 2]
        random.shuffle(indices)
        traces = [trace1, trace2, trace3]
        
        outer_idx, middle_idx, inner_idx = indices
        outer_trace = traces[outer_idx]
        middle_trace = traces[middle_idx]
        inner_trace = traces[inner_idx]
        
        # Prefer to have outer trace splittable; if not, but middle is, swap them
        if len(outer_trace) < 2 and len(middle_trace) >= 2:
            outer_trace, middle_trace = middle_trace, outer_trace
            outer_idx, middle_idx = middle_idx, outer_idx
        
        # Build segments for outer and middle
        if len(outer_trace) >= 2:
            outer_split = random.randint(1, len(outer_trace) - 1)
            outer_before = outer_trace.iloc[:outer_split].copy()
            outer_after = outer_trace.iloc[outer_split:].copy()
        else:
            # Can't split outer, keep it as a single block
            outer_before = outer_trace.copy()
            outer_after = None
        
        if len(middle_trace) >= 2:
            middle_split = random.randint(1, len(middle_trace) - 1)
            middle_before = middle_trace.iloc[:middle_split].copy()
            middle_after = middle_trace.iloc[middle_split:].copy()
        else:
            middle_before = middle_trace.copy()
            middle_after = None
        
        inner_full = inner_trace.copy()
        
        segments_in_order = []
        
        # Decide interleaving pattern based on what can be split
        if len(outer_trace) >= 2 and len(middle_trace) >= 2:
            # Ideal case: outer_before -> middle_before -> inner_full -> middle_after -> outer_after
            segments_in_order = [outer_before, middle_before, inner_full, middle_after, outer_after]
        elif len(outer_trace) >= 2:
            # Only outer can be split: outer_before -> middle_full -> inner_full -> outer_after
            segments_in_order = [outer_before, middle_before, inner_full, outer_after]
        elif len(middle_trace) >= 2:
            # Only middle can be split: outer_full -> middle_before -> inner_full -> middle_after
            segments_in_order = [outer_before, middle_before, inner_full, middle_after]
        else:
            # No splits possible: just concatenate in this random order
            segments_in_order = [outer_before, middle_before, inner_full]
        
        placed_segments = []
        first_segment = True
        
        for seg in segments_in_order:
            if seg is None or seg.empty:
                continue
            seg, current_max_time = place_segment(
                seg,
                current_max_time,
                add_gap_before=not first_segment  # gap before every segment except the first
            )
            placed_segments.append(seg)
            first_segment = False
        
        if not placed_segments:
            continue
        
        interleaved_trace = pd.concat(placed_segments, ignore_index=True)
        interleaved_trace = interleaved_trace.sort_values('time:timestamp').reset_index(drop=True)
        all_final_traces.append(interleaved_trace)
    
    if not all_final_traces:
        return pd.DataFrame()
    
    # Concatenate all traces and sort by timestamp
    result = pd.concat(all_final_traces, ignore_index=True)
    result = result.sort_values(by='time:timestamp').reset_index(drop=True)
    
    return result


def calculate_inter_trace_gaps(adjusted_logs):
    """Calculate gaps between traces"""
    inter_trace_gaps = {}
    for idx, log in enumerate(adjusted_logs):
        trace_times = []
        for case_id in log['case:concept:name'].unique():
            case_df = log[log['case:concept:name'] == case_id]
            start = case_df['time:timestamp'].min()
            end = case_df['time:timestamp'].max()
            trace_times.append((start, end))
        trace_times.sort()
        gaps = []
        for i in range(1, len(trace_times)):
            prev_end = trace_times[i-1][1]
            curr_start = trace_times[i][0]
            gap = (curr_start - prev_end).total_seconds()
            if gap > 0:
                gaps.append(gap)
        inter_trace_gaps[idx] = gaps
    variances = {idx: pd.Series(gaps).var() for idx, gaps in inter_trace_gaps.items()}
    return inter_trace_gaps, variances


def fit_best_distribution(data):
    """Fit best statistical distribution to data"""
    distributions = [stats.norm, stats.lognorm, stats.expon, stats.gamma, stats.beta]
    best_fit = None
    best_p = -1
    best_params = None
    for dist in distributions:
        try:
            params = dist.fit(data)
            D, p = stats.kstest(data, dist.name, args=params)
            if p > best_p:
                best_p = p
                best_fit = dist
                best_params = params
        except Exception:
            continue
    return best_fit, best_params


def fit_reference_distribution(time_diffs, variances, adjusted_logs, variance_criteria):
    """Fit reference distribution based on variance criteria"""
    reference_log_idx = max(variances, key=variances.get) if variance_criteria == 'max' else min(variances, key=variances.get)
    reference_log = adjusted_logs[reference_log_idx]
    time_diff_list = time_diffs[reference_log_idx]
    sorted_time_diffs = sorted(time_diff_list)
    best_dist, best_params = fit_best_distribution(sorted_time_diffs) if sorted_time_diffs else (None, None)
    return reference_log_idx, best_dist, best_params, sorted_time_diffs


def perform_first_round_shift(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs):
    """Perform first round of trace shifting"""
    updated_logs = []
    for log_idx, log in enumerate(adjusted_logs):
        if log_idx == reference_log_idx:
            updated_logs.append(log)
            continue
        new_times = []
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min()
        case_ids = case_starts.sort_values().index.tolist()
        last_end_time = None
        if sorted_time_diffs:
            try:
                lower_bound = max(0.0, float(np.percentile(sorted_time_diffs, 1)))
            except Exception:
                lower_bound = 0.0
            try:
                upper_bound = float(np.percentile(sorted_time_diffs, 99))
            except Exception:
                upper_bound = float(max(sorted_time_diffs))
            if not np.isfinite(upper_bound) or upper_bound <= 0:
                upper_bound = float(max(sorted_time_diffs)) if max(sorted_time_diffs) > 0 else 60.0
            upper_bound = float(min(upper_bound, 30 * 24 * 3600))
        else:
            lower_bound = 0.0
            upper_bound = 60.0
        for idx, case_id in enumerate(case_ids):
            case_df = log[log['case:concept:name'] == case_id].copy()
            case_df = case_df.sort_values(by='time:timestamp')
            if idx == 0:
                new_times.append(case_df)
                last_end_time = case_df['time:timestamp'].max()
            else:
                sampled_gap_seconds = None
                for _ in range(100):
                    try:
                        if best_dist is not None and best_params is not None:
                            candidate = float(best_dist.rvs(*best_params))
                        else:
                            candidate = float(random.choice(sorted_time_diffs)) if sorted_time_diffs else 1.0
                    except Exception:
                        candidate = 1.0
                    if not np.isfinite(candidate):
                        continue
                    candidate = max(lower_bound, min(candidate, upper_bound))
                    if candidate >= 0:
                        sampled_gap_seconds = candidate
                        break
                if sampled_gap_seconds is None:
                    sampled_gap_seconds = min(upper_bound, 60.0)
                orig_start = case_df['time:timestamp'].min()
                if last_end_time is not None:
                    new_start = last_end_time + pd.Timedelta(seconds=float(sampled_gap_seconds))
                else:
                    new_start = case_df['time:timestamp'].min()
                shift = new_start - orig_start
                case_df['time:timestamp'] = case_df['time:timestamp'] + shift
                last_end_time = case_df['time:timestamp'].max()
                new_times.append(case_df)
        updated_log = pd.concat(new_times, ignore_index=True)
        updated_logs.append(updated_log)
    try:
        first_round = pd.concat(updated_logs, ignore_index=True)
        out_dir = os.path.join('out', 'logs')
        os.makedirs(out_dir, exist_ok=True)
        first_round.sort_values(by=['time:timestamp']).to_csv(os.path.join(out_dir, 'after_first_round.csv'), index=False)
    except Exception:
        pass
    return updated_logs


def perform_gap_scaling_to_target(updated_logs, reference_log_idx):
    """Perform gap scaling to target"""
    try:
        ref_log = updated_logs[reference_log_idx]
        T_target = ref_log['time:timestamp'].max()
    except Exception:
        return updated_logs
    final_logs = []
    for log_idx, log in enumerate(updated_logs):
        if log_idx == reference_log_idx:
            final_logs.append(log)
            continue
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min().sort_values()
        ordered_case_ids = case_starts.index.tolist()
        if len(ordered_case_ids) <= 1:
            final_logs.append(log)
            continue
        trace_blocks = []
        for case_id in ordered_case_ids:
            case_df = log[log['case:concept:name'] == case_id].sort_values(by='time:timestamp').copy()
            trace_start = case_df['time:timestamp'].min()
            trace_end = case_df['time:timestamp'].max()
            trace_blocks.append((case_id, case_df, trace_start, trace_end))
        S_r = min(tb[2] for tb in trace_blocks)
        E_r = max(tb[3] for tb in trace_blocks)
        if pd.isna(S_r) or pd.isna(E_r) or T_target <= S_r:
            final_logs.append(log)
            continue
        gaps = []
        for i in range(1, len(trace_blocks)):
            prev_end = trace_blocks[i-1][3]
            next_start = trace_blocks[i][2]
            gap = (next_start - prev_end)
            gaps.append(gap)
        durations_sum = sum(((tb[3] - tb[2]) for tb in trace_blocks), pd.Timedelta(0))
        current_gaps_sum = sum(gaps, pd.Timedelta(0))
        desired_total_span = T_target - S_r
        new_gaps_sum = desired_total_span - durations_sum
        if new_gaps_sum < pd.Timedelta(0):
            new_gaps_sum = pd.Timedelta(0)
        if len(trace_blocks) == 1:
            final_logs.append(log)
            continue
        if current_gaps_sum > pd.Timedelta(0):
            scale = new_gaps_sum / current_gaps_sum if current_gaps_sum != pd.Timedelta(0) else 1.0
            new_gaps = [pd.to_timedelta(g * scale) for g in gaps]
        else:
            even_gap = new_gaps_sum / (len(trace_blocks) - 1) if (len(trace_blocks) - 1) > 0 else pd.Timedelta(0)
            new_gaps = [even_gap] * (len(trace_blocks) - 1)
        rebuilt = []
        first_case_df = trace_blocks[0][1].copy()
        rebuilt.append(first_case_df)
        rolling_end = first_case_df['time:timestamp'].max()
        for i in range(1, len(trace_blocks)):
            case_id, case_df, orig_start, orig_end = trace_blocks[i]
            desired_start = rolling_end + new_gaps[i-1]
            shift = desired_start - orig_start
            shifted_df = case_df.copy()
            shifted_df['time:timestamp'] = shifted_df['time:timestamp'] + shift
            rebuilt.append(shifted_df)
            rolling_end = shifted_df['time:timestamp'].max()
        final_logs.append(pd.concat(rebuilt, ignore_index=True))
    try:
        second_round = pd.concat(final_logs, ignore_index=True)
        out_dir = os.path.join('out', 'logs')
        os.makedirs(out_dir, exist_ok=True)
        second_round.sort_values(by=['time:timestamp']).to_csv(os.path.join(out_dir, 'after_second_round.csv'), index=False)
    except Exception:
        pass
    return final_logs


def shift_traces_by_distribution(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs):
    """Shift traces by distribution"""
    first_round_logs = perform_first_round_shift(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs)
    final_logs = perform_gap_scaling_to_target(first_round_logs, reference_log_idx)
    return final_logs


def compute_trace_interleaving_details(segmented_log):
    """
    Build per-trace overlap information.

    Returns a list of dictionaries with keys:
        - routine
        - case_id
        - start / end timestamps
        - overlap_routines (set of routines that overlap in time)
        - overlap_count (size of overlap_routines)
    """
    log_copy = segmented_log.copy()
    log_copy['time:timestamp'] = pd.to_datetime(log_copy['time:timestamp'])
    trace_intervals = []
    for routine in log_copy['routine_type'].unique():
        routine_log = log_copy[log_copy['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            trace_intervals.append({
                'routine': routine,
                'case_id': case_id,
                'start': trace['time:timestamp'].min(),
                'end': trace['time:timestamp'].max()
            })
    details = []
    for idx, t1 in enumerate(trace_intervals):
        overlaps = set()
        for jdx, t2 in enumerate(trace_intervals):
            if idx == jdx:
                continue
            if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                if t1['routine'] != t2['routine']:
                    overlaps.add(t2['routine'])
        details.append({
            'routine': t1['routine'],
            'case_id': t1['case_id'],
            'start': t1['start'],
            'end': t1['end'],
            'overlap_routines': overlaps,
            'overlap_count': len(overlaps)
        })
    return details


def count_trace_interleaving_cases(segmented_log):
    """Count trace interleaving cases"""
    segmented_log['time:timestamp'] = pd.to_datetime(segmented_log['time:timestamp'])
    trace_intervals = []
    for routine in segmented_log['routine_type'].unique():
        routine_log = segmented_log[segmented_log['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            start = trace['time:timestamp'].min()
            end = trace['time:timestamp'].max()
            trace_intervals.append({
                'routine': routine,
                'case_id': case_id,
                'start': start,
                'end': end
            })
    interleaved_with = dict()
    for i, t1 in enumerate(trace_intervals):
        overlaps = set()
        for j, t2 in enumerate(trace_intervals):
            if i == j:
                continue
            if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                if t1['routine'] != t2['routine']:
                    overlaps.add(t2['routine'])
        interleaved_with[(t1['routine'], t1['case_id'])] = overlaps
    count_0 = 0
    count_1 = 0
    count_2 = 0
    for overlap_set in interleaved_with.values():
        if len(overlap_set) == 0:
            count_0 += 1
        elif len(overlap_set) == 1:
            count_1 += 1
        elif len(overlap_set) >= 2:
            count_2 += 1
    print(f"\nNot interleaved: {count_0}")
    print(f"Interleaved with one: {count_1}")
    print(f"Interleaved with two or more: {count_2}\n")
    return {
        'not_interleaved': count_0,
        'interleaved_with_one': count_1,
        'interleaved_with_two_or_more': count_2
    }


def average_interleaving_per_log(segmented_log):
    """
    Compute the supervisor-requested metric: average # of interleaving routines per trace.

    Returns
    -------
    average : float
    std_dev : float
    per_trace_counts : List[int]
    """
    trace_details = compute_trace_interleaving_details(segmented_log)
    if not trace_details:
        return 0.0, 0.0, []
    counts = [detail['overlap_count'] for detail in trace_details]
    avg = float(np.mean(counts))
    std = float(np.std(counts, ddof=0))
    return avg, std, counts


def summarize_interleaving_samples(logs_by_class, doc_object=None, table_title="Average interleaving per class"):
    """
    logs_by_class: dict[label -> iterable of segmented logs].
        Each log item can be a pandas DataFrame or a CSV path.

    Returns (summary_df, samples_by_class, doc_object).
    """
    summary_rows = []
    samples_by_class = {}
    for class_name, log_items in logs_by_class.items():
        class_samples = []
        for log_item in log_items:
            if isinstance(log_item, pd.DataFrame):
                log_df = log_item
            elif isinstance(log_item, str):
                try:
                    log_df = pd.read_csv(log_item)
                except Exception as exc:
                    print(f"Skipping {log_item}: {exc}")
                    continue
            else:
                print(f"Unsupported log source for {class_name}: {type(log_item)}")
                continue
            avg, _, _ = average_interleaving_per_log(log_df)
            class_samples.append(avg)
        samples_by_class[class_name] = class_samples
        if class_samples:
            class_mean = float(np.mean(class_samples))
            class_std = float(np.std(class_samples, ddof=1)) if len(class_samples) > 1 else 0.0
        else:
            class_mean = class_std = 0.0
        summary_rows.append({
            'Class': class_name,
            '# Logs': len(class_samples),
            'Mean Avg Interleaving': class_mean,
            'Std Avg Interleaving': class_std
        })
    summary_df = pd.DataFrame(summary_rows)
    if not summary_df.empty:
        print(tabulate(summary_df, headers='keys', tablefmt='psql', showindex=False))
    if doc_object is not None and not summary_df.empty:
        doc_object.add_heading(table_title, level=2)
        table = doc_object.add_table(rows=1, cols=len(summary_df.columns))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for idx, col_name in enumerate(summary_df.columns):
            hdr[idx].text = str(col_name)
        for _, row in summary_df.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                if isinstance(value, float):
                    row_cells[idx].text = f"{value:.4f}"
                else:
                    row_cells[idx].text = str(value)
        doc_object.add_paragraph()
    return summary_df, samples_by_class, doc_object


def display_and_store_selected_routines(selected_routines, logs, document):
    """Display and store selected routines"""
    document.add_heading("Randomly Selected Logs", level=2)
    for num in selected_routines:
        text = f"Randomly select routine log {num} from log {logs[logs['routine_type'] == num]['log_number'].unique()[0]}"
        # display(HTML(styled_html))
        document.add_paragraph(text)
    return document


######################################################################################################################
####                                            Functions Called After process_random_logs                        ####
######################################################################################################################

def plot_interleaved_routines(final_log, numbers, doc_object, image_path="out/plots/interleaved_routines_plot.png", title="Interleaving of Selected Routines", duration='all', interleaved_only=False, clamp_to_threeway=False, plot_routine_level=False):
    """Plot interleaved routines"""
    if not os.path.exists(os.path.dirname(image_path)):
        os.makedirs(os.path.dirname(image_path))
        
    final_log['time:timestamp'] = pd.to_datetime(final_log['time:timestamp'])
    if duration in ['1h', '2h']:
        start_time = final_log['time:timestamp'].min()
        end_time = start_time + pd.Timedelta(hours=int(duration[0]))
        final_log = final_log[final_log['time:timestamp'] <= end_time]
    
    # Always plot individual traces
    trace_info, routines_to_plot = build_trace_intervals(final_log, numbers)
    
    if interleaved_only:
        interleaved_routines = set()
        interleaved_traces = []
        interleave_counts = defaultdict(int)
        for i in range(len(trace_info)):
            for j in range(i + 1, len(trace_info)):
                t1 = trace_info[i]
                t2 = trace_info[j]
                if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                    interleaved_traces.extend([t1, t2])
                    interleaved_routines.update([t1['routine'], t2['routine']])
                    pair = tuple(sorted([t1['routine'], t2['routine']]))
                    interleave_counts[pair] += 1
        unique_traces = {(t['routine'], t['case_id']): t for t in interleaved_traces}
        plot_traces = list(unique_traces.values())
        routines_to_plot = sorted(interleaved_routines)
    else:
        plot_traces = trace_info
    
    plot_traces_gantt(plot_traces, routines_to_plot, image_path, f"{title} ({duration}) - {len(routines_to_plot)} routines", duration, cutoff_time=None)
    doc_object.add_heading(f"{title} ({duration})", level=2)
    doc_object.add_picture(image_path, width=Inches(6.5))
    doc_object.add_paragraph(f"📊 Total Routines Plotted: {len(routines_to_plot)}")
    if interleaved_only:
        doc_object.add_paragraph("🔁 Interleaving Counts Between Routines:")
        for (r1, r2), count in interleave_counts.items():
            doc_object.add_paragraph(f" - Routine {r1} & Routine {r2}: {count} overlapping traces")
    return doc_object


def build_trace_intervals(final_log, numbers):
    """Build trace intervals"""
    trace_info = []
    for routine in numbers:
        routine_log = final_log[final_log['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            start = trace['time:timestamp'].min()
            end = trace['time:timestamp'].max()
            trace_info.append({
                'routine': routine,
                'case_id': case_id,
                'start': start,
                'end': end
            })
    routines_to_plot = sorted(set([t['routine'] for t in trace_info]))
    return trace_info, routines_to_plot


def plot_traces_gantt(trace_info, routines_to_plot, image_path, title, duration, cutoff_time=None):
    """Plot traces as Gantt chart"""
    plt.figure(figsize=(12, 5))
    routine_y_positions = {r: i for i, r in enumerate(routines_to_plot)}
    colors = plt.cm.get_cmap('tab10', len(routines_to_plot))
    for t in trace_info:
        y = routine_y_positions[t['routine']]
        line_start = t['start']
        line_end = t['end'] if cutoff_time is None else min(t['end'], cutoff_time)
        if cutoff_time is not None and line_start > cutoff_time:
            continue
        plt.hlines(
            y=y,
            xmin=line_start,
            xmax=line_end,
            colors=colors(routines_to_plot.index(t['routine'])),
            linewidth=6,
            label=f"Routine {t['routine']}"
        )
    handles, labels = plt.gca().get_legend_handles_labels()
    unique = dict(zip(labels, handles))
    plt.legend(unique.values(), unique.keys(), loc='upper center', bbox_to_anchor=(0.5, -0.15), ncol=len(unique))
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
    plt.gca().xaxis.set_major_locator(MaxNLocator(nbins=20))
    plt.yticks(list(routine_y_positions.values()), [f"Routine {r}" for r in routines_to_plot])
    plt.xlabel("Timestamp")
    plt.title(title)
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.tight_layout()
    plt.savefig(image_path)
    plt.close()
    return image_path


def plot_trace_interleaving_cases(segmented_log, image_path="out/plots/trace_interleaving_cases_bar.png", title="Trace Interleaving Cases"):
    """Plot trace interleaving cases as bar chart"""
    if not os.path.exists(os.path.dirname(image_path)):
        os.makedirs(os.path.dirname(image_path))
    counts = count_trace_interleaving_cases(segmented_log)
    categories = list(counts.keys())
    values = list(counts.values())
    total = sum(values)
    plt.figure(figsize=(8, 5))
    bars = plt.bar(categories, values, color=['#4e79a7', '#f28e2b', '#e15759'])
    plt.title(f"{title} (Total: {total})")
    plt.ylabel("Number of Traces")
    plt.xlabel("Interleaving Type")
    plt.grid(axis='y', linestyle='--', alpha=0.5)
    for bar, value in zip(bars, values):
        percent = (value / total * 100) if total > 0 else 0
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, f'{int(value)} ({percent:.0f}%)',
                 ha='center', va='bottom', fontsize=12, fontweight='bold')
    plt.tight_layout()
    plt.savefig(image_path)
    plt.close()


def update_avg_counts(trace_interleaved_counts):
    # Compute totals
    total = (
        trace_interleaved_counts['not_interleaved']
        + trace_interleaved_counts['interleaved_with_one']
        + trace_interleaved_counts['interleaved_with_two_or_more']
    )

    # Add averages
    trace_interleaved_counts['avg_not_interleaved'] = (
        trace_interleaved_counts['not_interleaved'] / total
    )
    trace_interleaved_counts['avg_interleaved'] = (
        (trace_interleaved_counts['interleaved_with_one']
        + trace_interleaved_counts['interleaved_with_two_or_more']) / total
    )

    return trace_interleaved_counts


######################################################################################################################
####                                            Evaluation and Export Functions                                   ####
######################################################################################################################

def jaccard_coefficient(set1, set2):
    """Calculate Jaccard coefficient"""
    return len(set1 & set2) / len(set1 | set2) if set1 | set2 else 0


def extract_transition_labels(petri_net):
    """Extract transition labels from Petri net"""
    return set([t.label for t in petri_net.transitions if t.label])


def extract_activity_names_and_log_ids_pre(activity_set):
    """Extract activity names and log IDs"""
    activity_names = []
    log_ids = set()
    
    for activity in activity_set:
        if '_' in activity:
            name, log_id = activity.rsplit('_', 1)
            activity_names.append(name)
            log_ids.add(log_id)
    
    return set(activity_names), log_ids


def extract_activity_names_and_log_ids(activity_set):
    """Extract activity names and log IDs (only when last part is numeric)."""
    activity_names = []
    log_ids = set()

    for activity in activity_set:
        if '_' in activity:
            name, last_part = activity.rsplit('_', 1)

            # Apply logic only if the last part is a number
            if last_part.isdigit():
                activity_names.append(name)
                log_ids.add(last_part)
            else:
                activity_names.append(activity)
        else:
            activity_names.append(activity)

    return set(activity_names), log_ids


def evaluate_clusters_with_jc(cluster_dict, models_folder="GT_Models", display_results=True, filter_cluster=False):
    """Evaluate clusters with Jaccard coefficient against all GT models (max JC)."""
    key_value = 1
    result_dict = {
        'Metrics': ['Cluster ID', "# Activities", "Unique Activities", 'JC'],
    }
    for cluster_id, cluster_activities in cluster_dict.items():
        activity_names, routine_ids = extract_activity_names_and_log_ids(cluster_activities)
        jc_scores = []
        jc_model_map = {}
        # Iterate over all ground truth models directly in models_folder
        if os.path.exists(models_folder):
            for file in os.listdir(models_folder):
                if not file.endswith('.pnml'):
                    continue
                pnml_file_path = os.path.join(models_folder, file)
                net, initial_marking, final_marking = pnml_importer.import_net(pnml_file_path)
                model_activities = extract_transition_labels(net)
                model_activities = set(
                    act for act in model_activities
                    if not (act.startswith('start') or act.startswith('end'))
                )
                jc = jaccard_coefficient(activity_names, model_activities)
                jc_scores.append(jc)
                jc_model_map[jc] = file
        max_jc = max(jc_scores) if jc_scores else 0
        best_jc_model = jc_model_map[max_jc] if max_jc in jc_model_map else "N/A"
        if display_results:
            print(f"Cluster {cluster_id}: Max JC = {max_jc:.4f}, Best GT Model: {best_jc_model}")
        result_dict[key_value] = [cluster_id, len(cluster_activities), len(activity_names), max_jc]
        key_value += 1
    return result_dict


def evaluate_clusters_with_jc_variant(cluster_dict, log_dataframe, logs_folder="logs", models_folder="GT_Models", display_results=True):
    """Evaluate clusters with Jaccard coefficient by extracting activities directly from cluster_dict and comparing with ground truth activities extracted from log_dataframe per routine type. Retains maximum JC across routine types."""
    key_value = 1
    result_dict = {
        'Metrics': ['Cluster ID', "# Activities", "Unique Activities", 'Trace Count', 'JC'],
    }
    for cluster_id, cluster_activities in cluster_dict.items():
        # Extract activity names directly from cluster_dict (no suffix removal needed, already clean)
        activity_names = set(cluster_activities) if not isinstance(cluster_activities, set) else cluster_activities
        
        # Find routine types that have at least one matching activity (clean log activities for comparison)
        # routine_types_in_cluster = []
        # for routine_type in log_dataframe['routine_type'].unique():
        #     routine_log = log_dataframe[log_dataframe['routine_type'] == routine_type]
        #     routine_activities_raw = set(routine_log['concept:name'].unique())
        #     # Clean activities by removing numeric suffix for comparison
        #     routine_activities_cleaned = set()
        #     for act in routine_activities_raw:
        #         if '_' in act:
        #             base, suffix = act.rsplit('_', 1)
        #             try:
        #                 int(suffix)
        #                 routine_activities_cleaned.add(base)
        #             except ValueError:
        #                 routine_activities_cleaned.add(act)
        #         else:
        #             routine_activities_cleaned.add(act)
        #     # Check if there's any overlap between cluster activities and cleaned routine activities
        #     if activity_names & routine_activities_cleaned:
        #         routine_types_in_cluster.append(routine_type)
        
        # # Filter dataframe for routine types in this cluster
        # subset = log_dataframe[log_dataframe['routine_type'].isin(routine_types_in_cluster)]
        # trace_count = int(subset['case:concept:name'].nunique()) if not subset.empty else 0
        
        # Extract ground truth activities for each routine type and calculate JC
        jc_by_routine = []  # Store (routine_type, jc, trace_count) tuples
        for routine_type in log_dataframe['routine_type'].unique():
            routine_log = log_dataframe[log_dataframe['routine_type'] == routine_type]
            # Extract unique activities from log for this routine type (ground truth)
            # No suffix removal needed as log activities are already clean
            gt_activities = set(routine_log['concept:name'].unique())
            
            # Calculate trace count for this routine type
            trace_count = int(routine_log['case:concept:name'].nunique())
            
            # Calculate JC between cluster activities and ground truth activities
            jc = jaccard_coefficient(activity_names, gt_activities)
            
            # Store routine type, JC, and trace count together
            jc_by_routine.append((routine_type, jc, trace_count))

            print(f"Routine type {routine_type}: JC = {jc:.4f}, Trace count = {trace_count} \n")
            print(f"Cluster activities: {activity_names} \n")
            print(f"Ground truth activities: {gt_activities} \n\n\n")
        
        # Find the routine type with maximum JC
        if jc_by_routine:
            best_routine_type, max_jc, trace_count = max(jc_by_routine, key=lambda x: x[1])
        else:
            best_routine_type, max_jc, trace_count = None, 0, 0
        
        if display_results:
            print(f"Cluster {cluster_id}: Max JC = {max_jc:.4f} from routine type {best_routine_type} (trace count: {trace_count})")
        
        result_dict[key_value] = [cluster_id, len(cluster_activities), len(activity_names), trace_count, max_jc]
        key_value += 1
    return result_dict


def append_averages_to_results(results_dict, selected_rotines):
    """Append averages to results"""
    selected_rotines = [str(routine) for routine in selected_rotines]
    routine_types_str = ', '.join(selected_rotines)
    simple_avg_label = f"Simple Average ({routine_types_str})"
    weighted_avg_label = f"Weighted Average ({routine_types_str})"
    headers = results_dict.get('Metrics', [])
    
    try:
        activities_idx = headers.index('# Activities')
    except ValueError:
        activities_idx = 1
    try:
        unique_activities_idx = headers.index('Unique Activities')
    except ValueError:
        unique_activities_idx = 2
    try:
        trace_count_idx = headers.index('Trace Count')
    except ValueError:
        trace_count_idx = 3
    try:
        jc_idx = headers.index('JC')
    except ValueError:
        jc_idx = len(headers) - 1

    simple_avg_row = [simple_avg_label]
    weighted_avg_row = [weighted_avg_label]
    total_activities = sum(results_dict[key][activities_idx] for key in results_dict if isinstance(key, int))
    total_unique_activities = sum(results_dict[key][unique_activities_idx] for key in results_dict if isinstance(key, int))
    total_trace_count = sum(results_dict[key][trace_count_idx] for key in results_dict if isinstance(key, int))
    simple_avg_row.append(str(total_activities))
    weighted_avg_row.append(str(total_activities))
    simple_avg_row.append(str(total_unique_activities))
    weighted_avg_row.append(str(total_unique_activities))
    simple_avg_row.append(str(total_trace_count))
    weighted_avg_row.append(str(total_trace_count))
    
    metric_values = [results_dict[key][jc_idx] for key in results_dict if isinstance(key, int)]
    simple_average = sum(metric_values) / len(metric_values) if metric_values else 0
    weighted_average = sum(
        results_dict[key][jc_idx] * results_dict[key][activities_idx]
        for key in results_dict if isinstance(key, int)
    ) / total_activities if total_activities else 0
    simple_avg_row.append(str(simple_average))
    weighted_avg_row.append(str(weighted_average))
    int_keys = [key for key in results_dict if isinstance(key, int)]
    next_index = (max(int_keys) + 1) if int_keys else 1
    results_dict[next_index] = simple_avg_row
    results_dict[next_index + 1] = weighted_avg_row
    return results_dict


def display_and_export_results_doc(results_dict, doc_object, table_title="Cluster Evaluation Summary"):
    """Display and export results to document"""
    result_df = pd.DataFrame(results_dict)
    result_df = result_df.set_index('Metrics')
    result_df = result_df.T
    print(tabulate(result_df, headers="keys", tablefmt="psql"))
    doc_object.add_heading(table_title, level=2)
    table = doc_object.add_table(rows=1, cols=len(result_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(result_df.columns):
        hdr_cells[i].text = str(col_name)
    for _, row in result_df.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = f"{cell_value:.4f}" if isinstance(cell_value, float) else str(cell_value)
    doc_object.add_paragraph()
    return doc_object


def collect_simple_average_rows_over_iterations(results_dicts):
    """Collect simple average rows over iterations"""
    rows = []
    for idx, results_dict in enumerate(results_dicts):
        # Count how many cluster IDs are present in this iteration
        num_clusters = 0
        for key, row in results_dict.items():
            if key == 'Metrics':
                continue
            # Treat rows whose first element is non-string as cluster rows
            if not isinstance(row[0], str):
                num_clusters += 1

        for key, row in results_dict.items():
            if isinstance(row[0], str) and row[0].startswith("Simple Average"):
                # Append iteration label, the Simple Average row, and number of clusters
                rows.append([f"Iter {idx+1}"] + row + [num_clusters])
                break
    if not rows:
        return []
    
    num_cols = len(rows[0])
    # Overall row: include placeholder labels in the first two columns
    mean_row = ["Overall Average", "Overall Mean"]
    for col in range(2, num_cols):
        col_values = []
        for row in rows:
            try:
                col_values.append(float(row[col]))
            except Exception:
                col_values.append(0)
        mean_val = sum(col_values) / len(col_values) if col_values else 0
        if all(isinstance(v, int) or (isinstance(v, float) and v.is_integer()) for v in col_values):
            mean_val = int(round(mean_val))
        mean_row.append(str(mean_val))
    rows.append(mean_row)
    return rows


def collect_iteration_metadata(all_results, all_iteration_metadata):
    """Collect iteration summary metadata for the new sheet with separate rows for each routine"""
    iteration_summary_rows = []
    
    for i, (results_dict, metadata) in enumerate(zip(all_results, all_iteration_metadata)):
        if metadata is None:
            continue
            
        # Extract data from metadata
        selected_routines = metadata.get('selected_routines', [])
        traces_before = metadata.get('traces_before_renumbering', {})
        traces_after = metadata.get('traces_after_renumbering', {})
        interleaving_counts = metadata.get('interleaving_counts', {})
        
        # Create summary row
        num_routines = len(selected_routines)
        
        # Format interleaving information (same for all rows in this iteration)
        total_interleaving = sum(interleaving_counts.values()) if interleaving_counts else 0
        interleaving_str = f"Total: {total_interleaving}"
        if interleaving_counts:
            interleaving_str += f" (Not interleaved: {interleaving_counts.get('not_interleaved', 0)}, "
            interleaving_str += f"With 1: {interleaving_counts.get('interleaved_with_one', 0)}, "
            interleaving_str += f"With 2+: {interleaving_counts.get('interleaved_with_two_or_more', 0)})"
        
        # Create a row for each routine
        for routine in selected_routines:
            iteration_summary_rows.append([
                f"Iteration {i+1}",
                num_routines,
                str(routine),
                f"R{routine}: {traces_before.get(routine, 0)}",
                f"R{routine}: {traces_after.get(routine, 0)}",
                interleaving_str
            ])
        
        # Add empty row between iterations (except for the last iteration)
        if i < len(all_iteration_metadata) - 1:
            iteration_summary_rows.append([
                '', '', '', '', '', ''
            ])
    
    return iteration_summary_rows


def display_and_export_results_xlx(all_results, all_iteration_metadata=None, output_path="Transformed_Logs_and_Results/Our", filename="all_iterations_full_results.xlsx", include_metadata=True):
    """Display and export results to Excel with optional metadata sheet"""
    # Build all iteration results as a list of rows
    all_rows = []
    for i, results_dict in enumerate(all_results):
        df = pd.DataFrame(results_dict)
        df = df.set_index('Metrics').T.reset_index()
        # Remove both 'Iteration' and 'index' columns if they exist
        for col in ['Iteration', 'index']:
            if col in df.columns:
                df = df.drop(columns=[col])
        # Add a label row (first cell is label, rest are empty)
        label_row = [f"Iteration {i+1}"] + [''] * (df.shape[1] - 1)
        all_rows.append(label_row)
        # Add column names as a row for this iteration
        all_rows.append(list(df.columns))
        # Add the iteration's results
        all_rows.extend(df.values.tolist())
        # Add a blank row for separation
        all_rows.append([''] * df.shape[1])
    
    # Remove the last blank row if present
    if all_rows and all(all_rows[-1][j] == '' for j in range(len(all_rows[-1]))):
        all_rows = all_rows[:-1]
    
    # Create DataFrame for all iterations
    iteration_results_df = pd.DataFrame(all_rows)
    
    # Create summary DataFrame
    summary_rows = collect_simple_average_rows_over_iterations(all_results)
    summary_df = pd.DataFrame(
        summary_rows,
        columns=["Iteration", "Label", "# Activities", "Unique Activities", "Trace Count", "JC", "Num Clusters"]
    )
    
    # Create sheets to write
    sheets_to_write = {
        "All Iterations": iteration_results_df,
        "Summary": summary_df
    }
    
    # Add iteration metadata sheet if requested and metadata is provided
    if include_metadata and all_iteration_metadata is not None:
        iteration_metadata_rows = collect_iteration_metadata(all_results, all_iteration_metadata)
        metadata_df = pd.DataFrame(
            iteration_metadata_rows,
            columns=["Iteration", "Num Selected Routines", "Selected Routines", "Traces Before Renumbering", "Traces After Renumbering", "Interleaving Counts"]
        )
        sheets_to_write["Iteration Metadata"] = metadata_df
    
    # Write all sheets to Excel
    with pd.ExcelWriter(f"{output_path}/{filename}") as writer:
        for sheet_name, df in sheets_to_write.items():
            if sheet_name == "All Iterations":
                df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
            else:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    sheet_names = list(sheets_to_write.keys())
    print(f"Results exported to {filename} with sheets: {', '.join(sheet_names)}")


def run_multiple_iterations_with_metadata(logs, document, num_iterations=5, target_datetime_str="2024-06-07 21:34:25.493914+00:00", variance_criteria='max'):
    """Run multiple iterations and collect metadata for comprehensive Excel export"""
    all_results = []
    all_iteration_metadata = []
    
    for iteration in range(num_iterations):
        print(f"\n=== Starting Iteration {iteration + 1} ===")
        
        # Process logs for this iteration
        result = process_random_logs(
            logs=logs,
            document=document,
            target_datetime_str=target_datetime_str,
            variance_criteria=variance_criteria
        )
        
        segment_log, unsegment_log, selected_routines, document, iteration_metadata = result
        
        if segment_log is not None:
            # For now, create a dummy results_dict
            results_dict = {
                'Metrics': ['Cluster ID', "# Activities", "Unique Activities", 'Trace Count', 'JC'],
                1: [1, 10, 8, iteration_metadata['traces_after_renumbering'].get(selected_routines[0], 0), 0.85]
            }
            
            all_results.append(results_dict)
            all_iteration_metadata.append(iteration_metadata)
            
            print(f"Iteration {iteration + 1} completed successfully")
        else:
            print(f"Iteration {iteration + 1} failed")
            all_iteration_metadata.append(None)
    
    # Export all results with metadata
    display_and_export_results_xlx(
        all_results=all_results,
        all_iteration_metadata=all_iteration_metadata,
        output_path="Transformed_Logs_and_Results/Our",
        filename="all_iterations_with_metadata.xlsx"
    )
    
    return all_results, all_iteration_metadata


######################################################################################################################
####                                            Post-Export Analysis & Plots                                      ####
######################################################################################################################

def plot_jc_boxplot_by_non_interleaving_bins(results_path, bins=None, output_path="out/plots/jc_boxplot_by_non_interleaving_bins.png"):
    """Read results (xlsx multi-sheet or csv), compute non-interleaving % per iteration, bin, and plot JC boxplots.

    Inputs:
    - results_path: Path to exported results. Supports the xlsx generated by display_and_export_results_xlx.
                    For csv, expects an Iteration Metadata-like CSV with columns including 'Iteration' and
                    'Interleaving Counts'.
    - bins: Optional list of bin edges (percentages). Default is [0,10,20,30,40,50,60,70,80,90,100].
    - output_path: Where to save the resulting boxplot image.

    Behavior:
    - For xlsx: reads 'Summary' and 'Iteration Metadata' sheets.
    - For csv: reads a single CSV assumed to mirror the 'Iteration Metadata' sheet structure.
    - Computes per-iteration non-interleaving percentage and aligns with Simple Average JC for that iteration.
    - Bins iterations by non-interleaving percentage and plots boxplots of JC per bin.
    """
    if bins is None:
        bins = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100]

    # Ensure output directory exists
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    # Helper to parse interleaving counts string like:
    # "Total: 123 (Not interleaved: 45, With 1: 60, With 2+: 18)"
    def _parse_interleaving_counts(counts_str):
        if not isinstance(counts_str, str) or not counts_str:
            return None, None, None, None
        total_match = re.search(r"Total:\s*(\d+)", counts_str)
        not_inter_match = re.search(r"Not\s+interleaved:\s*(\d+)", counts_str)
        with1_match = re.search(r"With\s+1:\s*(\d+)", counts_str)
        with2_match = re.search(r"With\s*2\+?:\s*(\d+)", counts_str)
        try:
            total = int(total_match.group(1)) if total_match else None
            n0 = int(not_inter_match.group(1)) if not_inter_match else None
            n1 = int(with1_match.group(1)) if with1_match else None
            n2 = int(with2_match.group(1)) if with2_match else None
        except Exception:
            return None, None, None, None
        return total, n0, n1, n2

    def _iter_label_to_int(label):
        try:
            return int(str(label).split()[-1])
        except Exception:
            return None

    def _load_xlsx(results_path):
        summary_df = pd.read_excel(results_path, sheet_name="Summary")
        metadata_df = pd.read_excel(results_path, sheet_name="Iteration Metadata")
        # Per-iteration Simple Average JC
        summary_df = summary_df.copy()
        summary_df["Iteration"] = summary_df["Iteration"].astype(str)
        per_iter_summary = summary_df[summary_df["Iteration"].str.startswith("Iter ")]
        per_iter_simple_avg = per_iter_summary[per_iter_summary["Label"].astype(str).str.startswith("Simple Average")]
        if per_iter_simple_avg.empty:
            per_iter_simple_avg = per_iter_summary
        jc_by_iter = per_iter_simple_avg.assign(iter_index=per_iter_simple_avg["Iteration"].map(_iter_label_to_int))[["iter_index", "JC"]].dropna()
        # Interleaving % per iteration
        metadata_df = metadata_df.copy()
        metadata_df["Iteration"] = metadata_df["Iteration"].astype(str)
        metadata_df = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")]
        def first_counts(group):
            vals = group["Interleaving Counts"].dropna().astype(str)
            for v in vals:
                if v.strip():
                    return v
            return None
        counts_by_iter_label = metadata_df.groupby("Iteration").apply(first_counts).reset_index(name="counts")
        counts_by_iter = counts_by_iter_label.assign(iter_index=counts_by_iter_label["Iteration"].map(_iter_label_to_int))
        parsed = counts_by_iter["counts"].apply(_parse_interleaving_counts)
        counts_by_iter = counts_by_iter.join(pd.DataFrame(parsed.tolist(), columns=["total", "not_inter", "with1", "with2"]))
        counts_by_iter["non_interleaving_pct"] = counts_by_iter.apply(
            lambda r: (r["not_inter"] / (r["not_inter"] + r["with1"] + r["with2"]) * 100.0) if all(
                isinstance(r[k], (int, float)) and r[k] is not None for k in ["not_inter", "with1", "with2"]
            ) and (r["not_inter"] + r["with1"] + r["with2"]) > 0 else None,
            axis=1
        )
        return pd.merge(jc_by_iter, counts_by_iter[["iter_index", "non_interleaving_pct"]], on="iter_index", how="inner")

    def _load_csv(results_path):
        metadata_df = pd.read_csv(results_path)
        required_cols = {"Iteration", "Interleaving Counts"}
        if not required_cols.issubset(set(metadata_df.columns)):
            raise ValueError("CSV must contain at least 'Iteration' and 'Interleaving Counts' columns.")
        metadata_df = metadata_df.copy()
        metadata_df["Iteration"] = metadata_df["Iteration"].astype(str)
        metadata_df = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")]
        counts_by_iter_label = metadata_df.groupby("Iteration")["Interleaving Counts"].apply(lambda s: next((v for v in s.astype(str) if v.strip()), None)).reset_index(name="counts")
        counts_by_iter = counts_by_iter_label.assign(iter_index=counts_by_iter_label["Iteration"].map(_iter_label_to_int))
        parsed = counts_by_iter["counts"].apply(_parse_interleaving_counts)
        counts_by_iter = counts_by_iter.join(pd.DataFrame(parsed.tolist(), columns=["total", "not_inter", "with1", "with2"]))
        counts_by_iter["non_interleaving_pct"] = counts_by_iter.apply(
            lambda r: (r["not_inter"] / (r["not_inter"] + r["with1"] + r["with2"]) * 100.0) if all(
                isinstance(r[k], (int, float)) and r[k] is not None for k in ["not_inter", "with1", "with2"]
            ) and (r["not_inter"] + r["with1"] + r["with2"]) > 0 else None,
            axis=1
        )
        if "JC" in metadata_df.columns:
            jc_by_iter = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")].copy()
            jc_by_iter = jc_by_iter.assign(iter_index=jc_by_iter["Iteration"].map(_iter_label_to_int))
            jc_by_iter = jc_by_iter.groupby("iter_index")["JC"].mean().reset_index()
            return pd.merge(jc_by_iter, counts_by_iter[["iter_index", "non_interleaving_pct"]], on="iter_index", how="inner")
        else:
            raise ValueError("CSV mode requires either the xlsx 'Summary' sheet or a 'JC' column in the CSV.")

    # Load data depending on extension
    if results_path.lower().endswith(".xlsx"):
        df = _load_xlsx(results_path)
    elif results_path.lower().endswith(".csv"):
        df = _load_csv(results_path)
    else:
        raise ValueError("Unsupported file type. Provide a .xlsx (recommended) or a .csv with required columns.")

    # Drop rows with missing values
    df = df.dropna(subset=["JC", "non_interleaving_pct"]).copy()

    if df.empty:
        raise ValueError("No valid data to plot after merging JC and non-interleaving percentages.")

    # Bin by non-interleaving percentage
    # Labels like '10-20', '20-30', ...
    bin_labels = [f"{int(bins[i])}-{int(bins[i+1])}" for i in range(len(bins)-1)]
    df["bin"] = pd.cut(df["non_interleaving_pct"], bins=bins, labels=bin_labels, include_lowest=True, right=False)

    # Prepare data for boxplot: list of JC arrays per bin in the order of bin_labels
    # Ensure every bin appears on x-axis even if empty by inserting [np.nan]
    data_per_bin = []
    xticklabels = bin_labels[:]
    for label in bin_labels:
        subset = df[df["bin"] == label]["JC"].astype(float)
        if subset.empty:
            data_per_bin.append([np.nan])
        else:
            data_per_bin.append(subset.values)

    plt.figure(figsize=(12, 6))
    bp = plt.boxplot(data_per_bin, labels=xticklabels, showfliers=True)
    plt.xlabel("Non-interleaving % bins")
    plt.ylabel("Simple Average JC per iteration")
    plt.title("JC distribution by non-interleaving percentage bins")
    plt.grid(axis='y', linestyle='--', alpha=0.5)

    # Add bin counts above x-ticks
    # Count true (non-NaN) observations per bin
    counts = []
    for label in xticklabels:
        c = int(df[df["bin"] == label]["JC"].notna().sum())
        counts.append(c)
    ymin, ymax = plt.ylim()
    # Position text slightly above upper whisker line
    y_text = ymax - (ymax - ymin) * 0.02
    for idx, (label, c) in enumerate(zip(xticklabels, counts), start=1):
        plt.text(idx, y_text, f"n={c}", ha='center', va='bottom', fontsize=9)

    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()
    print(f"Saved boxplot to {output_path}")



def read_simone_routines(routines_path):
    all_routines = []
    # Open the text file and read line by line
    with open(routines_path, "r") as file:
        for line in file:
            # Split the line into parts (splits by spaces by default)
            parts = line.strip().split()
            
            # Exclude the last item (supporting number) and store the routine list
            routine = parts[:-1]  # This excludes the last element
            all_routines.append(routine)  # Add the routine to the main list
    
    return all_routines


def read_dumas_routines(routines_path):
    all_routines = []
    # Open the text file and read line by line
    with open(routines_path, "r") as file:
        for line in file:
            # Use regex to extract the pattern between square brackets
            match = re.search(r"\[(.*?)\]", line)
            
            if match:
                # Split the pattern by commas to create a list of events
                routine = [event.strip() for event in match.group(1).split(',')]
                # Add the routine to the main list
                all_routines.append(routine)
    
    return all_routines


def calculate_average_jc_across_logs(results_base_path="out/results", technique="Our", noise_level="0.1", num_logs=100):
    """
    Calculate average JC across all logs by reading average rows from individual log result files.
    
    Parameters:
    - results_base_path: Base path where results are stored
    - noise_level: Noise level folder (e.g., "0.1")
    - num_logs: Number of logs to process (default 100)
    
    Returns:
    - Dictionary with statistics about JC values across all logs
    """
    jc_values = []
    log_stats = []
    
    print(f"Processing {num_logs} logs from {results_base_path}/Noise_{noise_level}/")
    
    for log_num in range(1, num_logs + 1):
        # Try both CSV and Excel formats
        csv_path = f"{results_base_path}/Noise_{noise_level}/{technique}_results_log{log_num}.csv"
        xlsx_path = f"{results_base_path}/Noise_{noise_level}/{technique}_results_log{log_num}.xlsx"
        
        file_path = None
        file_type = None
        
        # Check which file exists
        if os.path.exists(csv_path):
            file_path = csv_path
            file_type = 'csv'
        elif os.path.exists(xlsx_path):
            file_path = xlsx_path
            file_type = 'excel'
        else:
            print(f"Log {log_num}: No result file found (tried both .csv and .xlsx)")
            continue
        
        try:
            if file_type == 'csv':
                # Read CSV file
                df = pd.read_csv(file_path)
            else:
                # Read Excel file - try to find the summary sheet
                try:
                    # First try to read the 'Summary' sheet
                    df = pd.read_excel(file_path, sheet_name='Summary')
                except:
                    # If no Summary sheet, read the first sheet
                    df = pd.read_excel(file_path, sheet_name=0)
            
            print(f"Log {log_num}: Reading {file_type} file - {file_path}")
            print(f"Log {log_num}: Columns found: {list(df.columns)}")
            print(f"Log {log_num}: Shape: {df.shape}")
            
            # Look for JC column (case insensitive)
            jc_column = None
            for col in df.columns:
                if 'jc' in col.lower():
                    jc_column = col
                    break
            
            if jc_column is None:
                print(f"Log {log_num}: No 'JC' column found. Available columns: {list(df.columns)}")
                continue
            
            print(f"Log {log_num}: Using JC column: '{jc_column}'")
            
            # Strategy 1: Look for the LAST row that contains "Average" (this should be the final average)
            text_columns = df.select_dtypes(include=['object']).columns
            avg_row = None
            
            if len(text_columns) > 0:
                # Find all rows containing "Average"
                avg_mask = df[text_columns].astype(str).apply(
                    lambda x: x.str.contains('Average', case=False, na=False)
                ).any(axis=1)
                avg_rows = df[avg_mask]
                
                if not avg_rows.empty:
                    # Take the LAST average row (should be the final calculated average)
                    avg_row = avg_rows.iloc[-1]
                    print(f"Log {log_num}: Found {len(avg_rows)} average rows, using the LAST one")
                else:
                    # Strategy 2: Look for specific patterns in the last few rows
                    last_5_rows = df.tail(5)
                    for pattern in ['Simple Average', 'Weighted Average', 'Overall Average', 'Mean']:
                        pattern_mask = last_5_rows[text_columns].astype(str).apply(
                            lambda x: x.str.contains(pattern, case=False, na=False)
                        ).any(axis=1)
                        if pattern_mask.any():
                            avg_row = last_5_rows[pattern_mask].iloc[-1]  # Take the last matching row
                            print(f"Log {log_num}: Found pattern '{pattern}' in last 5 rows")
                            break
                    
                    # Strategy 3: If no pattern found, use the last row (assuming it's the average)
                    if avg_row is None:
                        avg_row = df.iloc[-1]
                        print(f"Log {log_num}: Using the last row as average")
            
            if avg_row is not None:
                # Get the single JC value from the average row
                jc_value = pd.to_numeric(avg_row[jc_column], errors='coerce')
                
                if not pd.isna(jc_value):
                    jc_values.append(jc_value)
                    log_stats.append({
                        'log_number': log_num,
                        'avg_jc': jc_value,
                        'file_type': file_type
                    })
                    print(f"Log {log_num}: Average JC = {jc_value:.4f} (single value from last average row)")
                else:
                    print(f"Log {log_num}: No valid JC value found in average row")
                    print(f"Log {log_num}: JC value in average row: {avg_row[jc_column]}")
            else:
                print(f"Log {log_num}: No average row found")
                # Debug: show last few rows
                print(f"Log {log_num}: Last 3 rows:")
                print(df.tail(3).to_string())
                
        except FileNotFoundError:
            print(f"Log {log_num}: File not found - {file_path}")
        except Exception as e:
            print(f"Log {log_num}: Error reading file - {str(e)}")
            import traceback
            traceback.print_exc()
    
    if not jc_values:
        print("No valid JC values found across any logs!")
        return None
    
    # Calculate overall statistics
    jc_array = np.array(jc_values)
    overall_stats = {
        'total_logs_processed': len(jc_values),
        'overall_mean_jc': np.mean(jc_array),
        'overall_std_jc': np.std(jc_array),
        'overall_min_jc': np.min(jc_array),
        'overall_max_jc': np.max(jc_array),
        'overall_median_jc': np.median(jc_array),
        'individual_log_stats': log_stats,
        'all_jc_values': jc_values
    }
    
    print(f"\n=== OVERALL STATISTICS ===")
    print(f"Total logs processed: {overall_stats['total_logs_processed']}")
    print(f"Overall Mean JC: {overall_stats['overall_mean_jc']:.4f}")
    print(f"Overall Std JC: {overall_stats['overall_std_jc']:.4f}")
    print(f"Overall Min JC: {overall_stats['overall_min_jc']:.4f}")
    print(f"Overall Max JC: {overall_stats['overall_max_jc']:.4f}")
    print(f"Overall Median JC: {overall_stats['overall_median_jc']:.4f}")
    
    return overall_stats


def export_overall_jc_summary(overall_stats, output_path="out/results", noise_level="0.1"):
    """
    Export overall JC summary to Excel file.
    
    Parameters:
    - overall_stats: Dictionary returned from calculate_average_jc_across_logs
    - output_path: Base output path
    - noise_level: Noise level for filename
    """
    if overall_stats is None:
        print("No statistics to export!")
        return
    
    # Create output directory
    os.makedirs(output_path, exist_ok=True)
    
    # Create summary DataFrame
    summary_data = {
        'Metric': [
            'Total Logs Processed',
            'Overall Mean JC',
            'Overall Std JC', 
            'Overall Min JC',
            'Overall Max JC',
            'Overall Median JC'
        ],
        'Value': [
            overall_stats['total_logs_processed'],
            f"{overall_stats['overall_mean_jc']:.4f}",
            f"{overall_stats['overall_std_jc']:.4f}",
            f"{overall_stats['overall_min_jc']:.4f}",
            f"{overall_stats['overall_max_jc']:.4f}",
            f"{overall_stats['overall_median_jc']:.4f}"
        ]
    }
    
    summary_df = pd.DataFrame(summary_data)
    
    # Create individual log statistics DataFrame
    individual_df = pd.DataFrame(overall_stats['individual_log_stats'])
    
    # Export to Excel with multiple sheets
    output_file = f"{output_path}/Overall_JC_Summary_Noise_{noise_level}.xlsx"
    
    with pd.ExcelWriter(output_file) as writer:
        summary_df.to_excel(writer, sheet_name='Overall Summary', index=False)
        individual_df.to_excel(writer, sheet_name='Individual Log Stats', index=False)
    
    print(f"Overall JC summary exported to: {output_file}")
    
    return output_file


def plot_jc_distribution_across_logs(overall_stats, output_path="out/plots", noise_level="0.1"):
    """
    Create plots showing JC distribution across all logs.
    
    Parameters:
    - overall_stats: Dictionary returned from calculate_average_jc_across_logs
    - output_path: Path to save plots
    - noise_level: Noise level for filename
    """
    if overall_stats is None:
        print("No statistics to plot!")
        return
    
    # Create output directory
    os.makedirs(output_path, exist_ok=True)
    
    jc_values = overall_stats['all_jc_values']
    
    # Create histogram
    plt.figure(figsize=(10, 6))
    plt.hist(jc_values, bins=20, alpha=0.7, color='skyblue', edgecolor='black')
    plt.axvline(overall_stats['overall_mean_jc'], color='red', linestyle='--', 
                label=f"Mean: {overall_stats['overall_mean_jc']:.4f}")
    plt.axvline(overall_stats['overall_median_jc'], color='green', linestyle='--', 
                label=f"Median: {overall_stats['overall_median_jc']:.4f}")
    plt.xlabel('Average JC per Log')
    plt.ylabel('Frequency')
    plt.title(f'Distribution of Average JC Across {overall_stats["total_logs_processed"]} Logs (Noise {noise_level})')
    plt.legend()
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    hist_path = f"{output_path}/jc_distribution_noise_{noise_level}.png"
    plt.savefig(hist_path)
    plt.close()
    
    # Create box plot
    plt.figure(figsize=(8, 6))
    plt.boxplot(jc_values, vert=True)
    plt.ylabel('Average JC per Log')
    plt.title(f'Box Plot of Average JC Across {overall_stats["total_logs_processed"]} Logs (Noise {noise_level})')
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    box_path = f"{output_path}/jc_boxplot_noise_{noise_level}.png"
    plt.savefig(box_path)
    plt.close()
    
    print(f"Plots saved to: {hist_path} and {box_path}")
    
    return hist_path, box_path


def analyze_jc_across_all_logs(results_base_path="out/results", technique="Our", noise_level="0.1", num_logs=100, 
                             export_results=True, create_plots=True):
    """
    Convenience function that combines all steps to analyze JC across all logs.
    
    This function:
    1. Reads average JC values from all individual log result files
    2. Calculates overall statistics
    3. Exports summary to Excel (optional)
    4. Creates distribution plots (optional)
    
    Parameters:
    - results_base_path: Base path where results are stored
    - noise_level: Noise level folder (e.g., "0.1")
    - num_logs: Number of logs to process (default 100)
    - export_results: Whether to export results to Excel (default True)
    - create_plots: Whether to create distribution plots (default True)
    
    Returns:
    - Dictionary with overall statistics
    """
    print(f"=== ANALYZING JC ACROSS ALL LOGS ===")
    print(f"Base path: {results_base_path}")
    print(f"Noise level: {noise_level}")
    print(f"Number of logs: {num_logs}")
    print("=" * 50)
    
    # Step 1: Calculate average JC across all logs
    overall_stats = calculate_average_jc_across_logs(
        results_base_path=results_base_path,
        technique=technique,
        noise_level=noise_level,
        num_logs=num_logs
    )
    
    if overall_stats is None:
        print("No valid data found. Exiting.")
        return None
    
    # Step 2: Export results to Excel (if requested)
    if export_results:
        print("\n=== EXPORTING RESULTS ===")
        excel_file = export_overall_jc_summary(
            overall_stats=overall_stats,
            output_path=results_base_path,
            noise_level=noise_level
        )
        print(f"Excel summary saved to: {excel_file}")
    
    # Step 3: Create plots (if requested)
    if create_plots:
        print("\n=== CREATING PLOTS ===")
        plot_paths = plot_jc_distribution_across_logs(
            overall_stats=overall_stats,
            output_path="out/plots",
            noise_level=noise_level
        )
        if plot_paths:
            print(f"Plots created: {plot_paths}")
    
    print("\n=== ANALYSIS COMPLETE ===")
    print(f"Successfully processed {overall_stats['total_logs_processed']} logs")
    print(f"Overall average JC: {overall_stats['overall_mean_jc']:.4f}")
    
    return overall_stats


def create_discovered_routines_csv(base_path="Transformed_Logs_and_Results/arebmann/Transformed_Log_With_Noise_0.1", 
                                 num_logs=100, prediction_value=-1):
    """
    Read log variants from input directories and create corresponding CSV files 
    in the Discovered_Routines directory structure.
    
    Parameters:
    - base_path: Base path containing the log directories
    - num_logs: Number of logs to process (default 100)
    - prediction_value: Always set to 1 as specified
    
    Creates CSV files with format:
    prediction,bound
    1,<number_of_rows_in_input_file>
    """
    import glob
    
    print(f"=== CREATING DISCOVERED ROUTINES CSV FILES ===")
    print(f"Base path: {base_path}")
    print(f"Number of logs: {num_logs}")
    print(f"Prediction value: {prediction_value}")
    print("=" * 60)
    
    successful_logs = 0
    failed_logs = []
    
    for log_num in range(1, num_logs + 1):
        # Input directory path
        input_dir = f"{base_path}/log{log_num}"
        
        # Output directory path
        output_dir = f"{base_path}/Discovered_Routines/log{log_num}"
        
        try:
            # Check if input directory exists
            if not os.path.exists(input_dir):
                print(f"Log {log_num}: Input directory not found - {input_dir}")
                failed_logs.append(f"Log {log_num}: Input directory not found")
                continue
            
            # Find all CSV files in the input directory
            input_files = glob.glob(os.path.join(input_dir, '*.csv'))
            
            if not input_files:
                print(f"Log {log_num}: No files found in input directory")
                failed_logs.append(f"Log {log_num}: No files found")
                continue
            
            # Create output directory
            os.makedirs(output_dir, exist_ok=True)
            
            # Process each file individually to keep all variants
            files_processed = 0
            for file_path in input_files:
                try:
                    df = pd.read_csv(file_path)
                    rows_in_file = len(df)
                    print(f"Log {log_num}: File {os.path.basename(file_path)} has {rows_in_file} rows")
                    
                    # Create output filename based on input filename
                    input_filename = os.path.basename(file_path)
                    # Add _pred before .csv extension
                    output_filename = input_filename.replace('.csv', '_pred.csv')
                    output_file = f"{output_dir}/{output_filename}"
                    
                    # Create DataFrame with the required format
                    csv_data = {
                        'prediction': [prediction_value],
                        'bound': [rows_in_file]
                    }
                    df_output = pd.DataFrame(csv_data)
                    
                    # Save to CSV
                    df_output.to_csv(output_file, index=False)
                    
                    print(f"Log {log_num}: Created {output_file} with prediction={prediction_value}, bound={rows_in_file}")
                    files_processed += 1
                    
                except Exception as e:
                    print(f"Log {log_num}: Error processing {os.path.basename(file_path)} - {str(e)}")
                    continue
            
            if files_processed > 0:
                successful_logs += 1
            else:
                print(f"Log {log_num}: No files were successfully processed")
                failed_logs.append(f"Log {log_num}: No files processed")
            
        except Exception as e:
            print(f"Log {log_num}: Error processing - {str(e)}")
            failed_logs.append(f"Log {log_num}: {str(e)}")
            continue
    
    # Summary
    print(f"\n=== SUMMARY ===")
    print(f"Successfully processed: {successful_logs} logs")
    print(f"Failed: {len(failed_logs)} logs")
    
    if failed_logs:
        print(f"\nFailed logs:")
        for failure in failed_logs:
            print(f"  - {failure}")
    
    return {
        'successful_logs': successful_logs,
        'failed_logs': len(failed_logs),
        'failed_details': failed_logs
    }


def create_discovered_routines_csv_advanced(base_path="Transformed_Logs_and_Results/arebmann/Transformed_Log_With_Noise_0.1", 
                                          num_logs=100, prediction_value=1, 
                                          output_filename="discovered_routines.csv"):
    """
    Advanced version with more options for creating discovered routines CSV files.
    Only processes CSV files.
    
    Parameters:
    - base_path: Base path containing the log directories
    - num_logs: Number of logs to process (default 100)
    - prediction_value: Always set to -1 as specified
    - output_filename: Name of the output CSV file (fallback only)
    """
    import glob
    
    print(f"=== CREATING DISCOVERED ROUTINES CSV FILES (ADVANCED) ===")
    print(f"Base path: {base_path}")
    print(f"Number of logs: {num_logs}")
    print(f"Prediction value: {prediction_value}")
    print(f"Output filename: {output_filename}")
    print("=" * 70)
    
    successful_logs = 0
    failed_logs = []
    detailed_stats = []
    
    for log_num in range(1, num_logs + 1):
        # Input directory path
        input_dir = f"{base_path}/log{log_num}"
        
        # Output directory path
        output_dir = f"{base_path}/Discovered_Routines/log{log_num}"
        
        log_stats = {
            'log_number': log_num,
            'input_dir': input_dir,
            'output_dir': output_dir,
            'files_found': 0,
            'total_rows': 0,
            'success': False,
            'error': None
        }
        
        try:
            # Check if input directory exists
            if not os.path.exists(input_dir):
                error_msg = f"Input directory not found - {input_dir}"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
                detailed_stats.append(log_stats)
                continue
            
            # Find all CSV files in the input directory
            input_files = glob.glob(os.path.join(input_dir, '*.csv'))
            
            log_stats['files_found'] = len(input_files)
            
            if not input_files:
                error_msg = "No CSV files found in input directory"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
                detailed_stats.append(log_stats)
                continue
            
            # Create output directory
            os.makedirs(output_dir, exist_ok=True)
            
            # Process each file individually to keep all variants
            files_processed = 0
            file_details = []
            total_rows = 0
            
            for file_path in input_files:
                try:
                    df = pd.read_csv(file_path)
                    rows_in_file = len(df)
                    total_rows += rows_in_file
                    
                    # Create output filename based on input filename
                    input_filename = os.path.basename(file_path)
                    # Add _pred before .csv extension
                    output_filename = input_filename.replace('.csv', '_pred.csv')
                    output_file = f"{output_dir}/{output_filename}"
                    
                    # Create DataFrame with the required format
                    csv_data = {
                        'prediction': [prediction_value],
                        'bound': [rows_in_file]
                    }
                    df_output = pd.DataFrame(csv_data)
                    
                    # Save to CSV
                    df_output.to_csv(output_file, index=False)
                    
                    print(f"Log {log_num}: Created {output_file}")
                    print(f"  - File: {os.path.basename(file_path)}")
                    print(f"  - Rows: {rows_in_file}")
                    print(f"  - Prediction: {prediction_value}, Bound: {rows_in_file}")
                    
                    file_details.append({
                        'filename': os.path.basename(file_path),
                        'output_file': output_filename,
                        'rows': rows_in_file
                    })
                    
                    files_processed += 1
                    
                except Exception as e:
                    print(f"Log {log_num}: Error processing {os.path.basename(file_path)} - {str(e)}")
                    continue
            
            log_stats['total_rows'] = total_rows
            log_stats['file_details'] = file_details
            log_stats['files_processed'] = files_processed
            
            if files_processed > 0:
                log_stats['success'] = True
                successful_logs += 1
                print(f"Log {log_num}: Successfully processed {files_processed} files")
            else:
                error_msg = "No files were successfully processed"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
            
        except Exception as e:
            error_msg = f"Error processing - {str(e)}"
            print(f"Log {log_num}: {error_msg}")
            log_stats['error'] = error_msg
            failed_logs.append(f"Log {log_num}: {error_msg}")
        
        detailed_stats.append(log_stats)
    
    # Summary
    print(f"\n=== DETAILED SUMMARY ===")
    print(f"Successfully processed: {successful_logs} logs")
    print(f"Failed: {len(failed_logs)} logs")
    
    if failed_logs:
        print(f"\nFailed logs:")
        for failure in failed_logs:
            print(f"  - {failure}")
    
    # Create summary report
    summary_df = pd.DataFrame(detailed_stats)
    summary_file = f"{base_path}/discovered_routines_summary.xlsx"
    
    try:
        with pd.ExcelWriter(summary_file) as writer:
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        print(f"\nDetailed summary saved to: {summary_file}")
    except Exception as e:
        print(f"Warning: Could not save summary file - {str(e)}")
    
    return {
        'successful_logs': successful_logs,
        'failed_logs': len(failed_logs),
        'failed_details': failed_logs,
        'detailed_stats': detailed_stats,
        'summary_file': summary_file
    }


def remove_interleaving_sequential_routines(log_input, gap_seconds=1.0, output_path=None):
    """
    Remove interleaving from a log by sequencing traces: all traces of routine 1 first, 
    then all traces of routine 2. Timestamps are adjusted to maintain sequential order 
    while preserving inter-action gaps within each trace.
    
    Parameters:
    -----------
    log_input : str or pd.DataFrame
        Path to CSV file or DataFrame containing the interleaved log.
        Must have columns: 'time:timestamp', 'case:concept:name', 'routine_type'
    gap_seconds : float, default=1.0
        Gap in seconds to add between consecutive traces of the same routine type.
    output_path : str, optional
        If provided, save the de-interleaved log to this CSV path.
    
    Returns:
    --------
    pd.DataFrame
        De-interleaved log with all routine 1 traces first, then all routine 2 traces.
    
    Example:
    --------
    >>> deinterleaved_log = remove_interleaving_sequential_routines(
    ...     'interleaved_log.csv', 
    ...     gap_seconds=2.0,
    ...     output_path='deinterleaved_log.csv'
    ... )
    """
    # Load log if path provided
    if isinstance(log_input, str):
        log = pd.read_csv(log_input)
    elif isinstance(log_input, pd.DataFrame):
        log = log_input.copy()
    else:
        raise ValueError("log_input must be a file path (str) or DataFrame")
    
    # Ensure required columns exist
    required_columns = ['time:timestamp', 'case:concept:name', 'routine_type']
    missing_columns = [col for col in required_columns if col not in log.columns]
    if missing_columns:
        raise ValueError(f"Missing required columns: {missing_columns}")
    
    # Convert timestamp to datetime
    log['time:timestamp'] = pd.to_datetime(log['time:timestamp'])
    
    # Get unique routine types
    routine_types = sorted(log['routine_type'].unique())
    
    if len(routine_types) != 2:
        raise ValueError(f"Expected exactly 2 routine types, found {len(routine_types)}: {routine_types}")
    
    routine_1, routine_2 = routine_types[0], routine_types[1]
    
    print(f"Processing log with routine types: {routine_1} and {routine_2}")
    
    # Separate logs by routine type
    log_routine_1 = log[log['routine_type'] == routine_1].copy()
    log_routine_2 = log[log['routine_type'] == routine_2].copy()
    
    # Get all traces for each routine type
    def get_traces_sorted(log_df):
        """Get all traces sorted by their start time"""
        case_starts = log_df.groupby('case:concept:name')['time:timestamp'].min().sort_values()
        traces = []
        for case_id in case_starts.index:
            trace = log_df[log_df['case:concept:name'] == case_id].sort_values('time:timestamp').copy()
            traces.append((case_id, trace))
        return traces
    
    traces_routine_1 = get_traces_sorted(log_routine_1)
    traces_routine_2 = get_traces_sorted(log_routine_2)
    
    print(f"Routine {routine_1}: {len(traces_routine_1)} traces")
    print(f"Routine {routine_2}: {len(traces_routine_2)} traces")
    
    # Process traces sequentially, preserving inter-action gaps
    def process_traces_sequentially(traces, start_time=None):
        """
        Process traces sequentially, shifting timestamps to maintain order.
        Preserves inter-action gaps within each trace.
        """
        if not traces:
            return pd.DataFrame()
        
        processed_traces = []
        current_max_time = start_time
        
        for case_id, trace in traces:
            trace = trace.copy()
            
            if current_max_time is None:
                # First trace: use its original start time
                trace_start = trace['time:timestamp'].min()
                current_max_time = trace['time:timestamp'].max()
            else:
                # Subsequent traces: shift to start after previous trace ends + gap
                trace_start_original = trace['time:timestamp'].min()
                trace_end_original = trace['time:timestamp'].max()
                
                # Calculate the duration of the trace
                trace_duration = trace_end_original - trace_start_original
                
                # New start time: after previous trace ends + gap
                new_start_time = current_max_time + pd.Timedelta(seconds=gap_seconds)
                
                # Calculate shift needed
                shift = new_start_time - trace_start_original
                
                # Apply shift to all timestamps in trace (preserves inter-action gaps)
                trace['time:timestamp'] = trace['time:timestamp'] + shift
                
                # Update current_max_time to the end of this trace
                current_max_time = trace['time:timestamp'].max()
            
            processed_traces.append(trace)
        
        # Concatenate all processed traces
        if processed_traces:
            return pd.concat(processed_traces, ignore_index=True)
        else:
            return pd.DataFrame()
    
    # Process routine 1 traces first
    print(f"\nProcessing routine {routine_1} traces...")
    processed_routine_1 = process_traces_sequentially(traces_routine_1, start_time=None)
    
    # Get the end time of routine 1 to start routine 2 after it
    routine_1_end_time = None
    if not processed_routine_1.empty:
        routine_1_end_time = processed_routine_1['time:timestamp'].max()
        print(f"Routine {routine_1} ends at: {routine_1_end_time}")
    else:
        # If routine 1 is empty, use the minimum time from routine 2
        if not log_routine_2.empty:
            routine_1_end_time = log_routine_2['time:timestamp'].min() - pd.Timedelta(seconds=gap_seconds)
        else:
            routine_1_end_time = pd.Timestamp.now()
    
    # Process routine 2 traces starting after routine 1 ends
    print(f"\nProcessing routine {routine_2} traces...")
    processed_routine_2 = process_traces_sequentially(
        traces_routine_2, 
        start_time=routine_1_end_time + pd.Timedelta(seconds=gap_seconds)
    )
    
    if not processed_routine_2.empty:
        routine_2_end_time = processed_routine_2['time:timestamp'].max()
        print(f"Routine {routine_2} ends at: {routine_2_end_time}")
    
    # Combine both routine types
    if not processed_routine_1.empty and not processed_routine_2.empty:
        deinterleaved_log = pd.concat([processed_routine_1, processed_routine_2], ignore_index=True)
    elif not processed_routine_1.empty:
        deinterleaved_log = processed_routine_1
    elif not processed_routine_2.empty:
        deinterleaved_log = processed_routine_2
    else:
        raise ValueError("No traces found in either routine type")
    
    # Sort by timestamp
    deinterleaved_log = deinterleaved_log.sort_values('time:timestamp').reset_index(drop=True)
    
    # Verify no interleaving remains
    routine_1_max = deinterleaved_log[deinterleaved_log['routine_type'] == routine_1]['time:timestamp'].max()
    routine_2_min = deinterleaved_log[deinterleaved_log['routine_type'] == routine_2]['time:timestamp'].min()
    
    if not pd.isna(routine_1_max) and not pd.isna(routine_2_min):
        if routine_2_min < routine_1_max:
            print(f"Warning: Some interleaving may remain (routine {routine_2} starts before routine {routine_1} ends)")
        else:
            print(f"✓ De-interleaving successful: Routine {routine_2} starts after routine {routine_1} ends")
    
    # Save to file if output path provided
    if output_path:
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        deinterleaved_log.to_csv(output_path, index=False)
        print(f"\nDe-interleaved log saved to: {output_path}")
    
    return deinterleaved_log


# def normalize_activity_sequences(log_input, output_path=None, min_gap_seconds=1.0):
#     """
#     Normalize activity sequences within traces to match canonical sequences per routine type.
#     Preserves trace duration and ensures no trace overlaps with the next trace.
    
#     For each routine type:
#     1. Finds the most common activity sequence (canonical sequence) *using base labels* (stripping suffixes)
#     2. Reorders events in each trace to match the canonical base sequence
#     3. Adjusts timestamps to maintain correct order while preserving relative gaps
#     4. Preserves overall trace duration
#     5. Ensures next trace starts after current trace's max time
    
#     Parameters:
#     -----------
#     log_input : str or pd.DataFrame
#         Path to CSV file or DataFrame containing the log.
#         Must have columns: 'time:timestamp', 'case:concept:name', 'routine_type', 'concept:name'
#     output_path : str, optional
#         If provided, save the normalized log to this CSV path.
#     min_gap_seconds : float, default=1.0
#         Minimum gap in seconds between consecutive traces to prevent overlap.
    
#     Returns:
#     --------
#     pd.DataFrame
#         Log with normalized activity sequences per routine type.
#     """
#     # Load log if path provided
#     if isinstance(log_input, str):
#         log = pd.read_csv(log_input)
#     elif isinstance(log_input, pd.DataFrame):
#         log = log_input.copy()
#     else:
#         raise ValueError("log_input must be a file path (str) or DataFrame")
    
#     # Ensure required columns exist
#     required_columns = ['time:timestamp', 'case:concept:name', 'routine_type', 'concept:name']
#     missing_columns = [col for col in required_columns if col not in log.columns]
#     if missing_columns:
#         raise ValueError(f"Missing required columns: {missing_columns}")
    
#     # Convert timestamp to datetime
#     log['time:timestamp'] = pd.to_datetime(log['time:timestamp'])
    
#     print("=" * 60)
#     print("NORMALIZING ACTIVITY SEQUENCES")
#     print("=" * 60)
    
#     # Helper to get base label
#     def get_base_label(label):
#         if not isinstance(label, str):
#             return str(label)
#         if '_' in label:
#             parts = label.rsplit('_', 1)
#             if len(parts) == 2 and parts[1].isdigit():
#                 return parts[0]
#         return label

#     # Get all routine types
#     routine_types = sorted(log['routine_type'].unique())
#     print(f"Found {len(routine_types)} routine type(s): {routine_types}")
    
#     # Step 1: Find canonical sequences for each routine type (using base labels)
#     def find_canonical_sequence(log_df, routine_type):
#         """Find the most common activity sequence (canonical sequence on base labels) for a routine type"""
#         routine_log = log_df[log_df['routine_type'] == routine_type].copy()
        
#         # Get sequence for each trace
#         sequences = []
#         for case_id in routine_log['case:concept:name'].unique():
#             trace = routine_log[routine_log['case:concept:name'] == case_id].sort_values('time:timestamp')
#             # Use base labels for canonical sequence
#             sequence = tuple([get_base_label(act) for act in trace['concept:name'].tolist()])
#             sequences.append(sequence)
        
#         if not sequences:
#             return []
        
#         # Find most common sequence (mode)
#         from collections import Counter
#         sequence_counts = Counter(sequences)
#         canonical_sequence = sequence_counts.most_common(1)[0][0]
#         canonical_count = sequence_counts[canonical_sequence]
#         total_traces = len(sequences)
        
#         print(f"\nRoutine {routine_type}:")
#         print(f"  Total traces: {total_traces}")
#         print(f"  Canonical sequence (base): {list(canonical_sequence)}")
#         print(f"  Traces matching canonical: {canonical_count} ({canonical_count/total_traces*100:.1f}%)")
#         print(f"  Unique sequences (base): {len(sequence_counts)}")
        
#         return list(canonical_sequence)
    
#     canonical_sequences = {}
#     for routine_type in routine_types:
#         canonical_sequences[routine_type] = find_canonical_sequence(log, routine_type)
    
#     # Step 2: Normalize each trace to match canonical sequence
#     def normalize_trace(trace, canonical_sequence):
#         """
#         Normalize a trace to match canonical sequence while preserving duration and gaps.
#         Uses base labels for matching.
#         """
#         trace = trace.copy().sort_values('time:timestamp').reset_index(drop=True)
        
#         # Get original trace metadata
#         original_start = trace['time:timestamp'].min()
#         original_end = trace['time:timestamp'].max()
#         original_duration = original_end - original_start
        
#         # Get current sequence (base labels only for comparison)
#         current_sequence_base = [get_base_label(act) for act in trace['concept:name'].tolist()]
        
#         # If already matches canonical, return as is
#         if current_sequence_base == canonical_sequence:
#             return trace
        
#         # Calculate original gaps between consecutive events (preserve these)
#         original_gaps = []
#         original_times = trace['time:timestamp'].tolist()
#         for i in range(len(original_times) - 1):
#             gap = (original_times[i+1] - original_times[i]).total_seconds()
#             original_gaps.append(max(0.1, gap))  # Ensure minimum gap
        
#         # Build mapping: match events from trace to canonical positions
#         from collections import Counter, defaultdict
        
#         # Create availability lists for each event type in canonical
#         canonical_positions = defaultdict(list)
#         for pos, event in enumerate(canonical_sequence):
#             canonical_positions[event].append(pos)
        
#         # Greedy matching: for each event in trace, assign to earliest available canonical position
#         trace_mapping = []  # List of (trace_index, canonical_position, event_name)
#         used_positions = set()
        
#         for trace_idx, row in trace.iterrows():
#             event = row['concept:name']
#             base_event = get_base_label(event)
            
#             # Try to find a matching position in canonical sequence (using base_event)
#             if base_event in canonical_positions:
#                 available_positions = [p for p in canonical_positions[base_event] 
#                                      if p not in used_positions]
#                 if available_positions:
#                     # Use the earliest available position
#                     target_pos = min(available_positions)
#                     trace_mapping.append((trace_idx, target_pos, event))
#                     used_positions.add(target_pos)
#                     continue
            
#             # Event not in canonical or all positions used - place after canonical sequence
#             target_pos = len(canonical_sequence) + len([m for m in trace_mapping if m[1] >= len(canonical_sequence)])
#             trace_mapping.append((trace_idx, target_pos, event))
        
#         # Sort by canonical position to get new order
#         trace_mapping.sort(key=lambda x: (x[1], x[0]))  # Sort by canonical pos, then by original index
        
#         # Reorder trace according to mapping
#         reordered_indices = [idx for idx, _, _ in trace_mapping]
#         normalized_trace = trace.iloc[reordered_indices].copy().reset_index(drop=True)
        
#         # Assign new timestamps preserving original gaps and duration
#         new_timestamps = [original_start]
        
#         # Use original gaps, maintaining the same relative timing
#         for i in range(1, len(normalized_trace)):
#             if i-1 < len(original_gaps):
#                 gap_seconds = original_gaps[i-1]
#             else:
#                 # If we have more events than original gaps, use average gap
#                 avg_gap = np.mean(original_gaps) if original_gaps else 1.0
#                 gap_seconds = avg_gap
            
#             next_time = new_timestamps[-1] + pd.Timedelta(seconds=gap_seconds)
#             new_timestamps.append(next_time)
        
#         # Scale to preserve original duration
#         new_end = new_timestamps[-1]
#         new_duration = (new_end - original_start).total_seconds()
#         original_duration_seconds = original_duration.total_seconds()
        
#         if new_duration > 0 and original_duration_seconds > 0:
#             if new_duration > original_duration_seconds:
#                 # Scale down to fit within original duration
#                 scale_factor = original_duration_seconds / new_duration
#                 scaled_timestamps = [original_start]
#                 for i in range(1, len(new_timestamps)):
#                     gap = (new_timestamps[i] - new_timestamps[i-1]).total_seconds() * scale_factor
#                     scaled_timestamps.append(scaled_timestamps[-1] + pd.Timedelta(seconds=gap))
#                 new_timestamps = scaled_timestamps
#             # If new_duration <= original_duration, keep as is (preserves gaps better)
        
#         # Assign new timestamps
#         normalized_trace['time:timestamp'] = new_timestamps
        
#         return normalized_trace
    
#     # Step 3: Process all traces, ensuring no overlaps
#     normalized_traces = []
    
#     for routine_type in routine_types:
#         routine_log = log[log['routine_type'] == routine_type].copy()
#         canonical_sequence = canonical_sequences[routine_type]
        
#         if not canonical_sequence:
#             print(f"\nWarning: No canonical sequence found for routine {routine_type}, skipping...")
#             normalized_traces.append(routine_log)
#             continue
        
#         # Get all traces for this routine, sorted by start time
#         case_starts = routine_log.groupby('case:concept:name')['time:timestamp'].min().sort_values()
#         routine_traces = []
        
#         for case_id in case_starts.index:
#             trace = routine_log[routine_log['case:concept:name'] == case_id].copy()
#             routine_traces.append((case_id, trace))
        
#         print(f"\nNormalizing {len(routine_traces)} traces for routine {routine_type}...")
        
#         # Process traces sequentially, ensuring no overlap
#         current_max_time = None
#         normalized_routine_traces = []
        
#         for case_id, trace in routine_traces:
#             # Normalize the trace sequence
#             normalized_trace = normalize_trace(trace, canonical_sequence)
            
#             # Get trace timing
#             trace_start = normalized_trace['time:timestamp'].min()
#             trace_end = normalized_trace['time:timestamp'].max()
#             trace_duration = trace_end - trace_start
            
#             # Adjust start time if needed to prevent overlap
#             if current_max_time is not None:
#                 # Ensure this trace starts after previous trace ends
#                 required_start = current_max_time + pd.Timedelta(seconds=min_gap_seconds)
#                 if trace_start < required_start:
#                     # Shift entire trace forward
#                     shift = required_start - trace_start
#                     normalized_trace['time:timestamp'] = normalized_trace['time:timestamp'] + shift
#                     trace_start = normalized_trace['time:timestamp'].min()
#                     trace_end = normalized_trace['time:timestamp'].max()
            
#             # Update current_max_time
#             current_max_time = trace_end
            
#             normalized_routine_traces.append(normalized_trace)
        
#         # Combine all traces for this routine
#         if normalized_routine_traces:
#             normalized_traces.append(pd.concat(normalized_routine_traces, ignore_index=True))
    
#     # Combine all routine types
#     if normalized_traces:
#         normalized_log = pd.concat(normalized_traces, ignore_index=True)
#     else:
#         normalized_log = log.copy()
    
#     # Sort by timestamp
#     normalized_log = normalized_log.sort_values('time:timestamp').reset_index(drop=True)
    
#     # Statistics
#     print("\n" + "=" * 60)
#     print("NORMALIZATION SUMMARY")
#     print("=" * 60)
    
#     for routine_type in routine_types:
#         routine_log_original = log[log['routine_type'] == routine_type]
#         routine_log_normalized = normalized_log[normalized_log['routine_type'] == routine_type]
        
#         original_traces = routine_log_original['case:concept:name'].nunique()
#         normalized_traces = routine_log_normalized['case:concept:name'].nunique()
        
#         original_duration = (routine_log_original['time:timestamp'].max() - 
#                             routine_log_original['time:timestamp'].min()).total_seconds()
#         normalized_duration = (routine_log_normalized['time:timestamp'].max() - 
#                               routine_log_normalized['time:timestamp'].min()).total_seconds()
        
#         print(f"\nRoutine {routine_type}:")
#         print(f"  Traces: {original_traces} → {normalized_traces}")
#         print(f"  Total duration: {original_duration:.1f}s → {normalized_duration:.1f}s")
    
#     # Save to file if output path provided
#     if output_path:
#         os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
#         normalized_log.to_csv(output_path, index=False)
#         print(f"\n✓ Normalized log saved to: {output_path}")
    
#     return normalized_log


import os
import numpy as np
import pandas as pd
from collections import Counter, defaultdict, deque


# ============================================================
# Helper 1: Canonical Sequence Discovery (Full Activity Labels)
# ============================================================

def find_canonical_sequence(log_df, routine_type):
    """
    Finds the most common activity sequence (canonical sequence)
    for a given routine type.

    Important:
    ----------
    - Uses full activity labels exactly as they appear
    - Preserves duplicates (e.g., b appears twice)
    """

    routine_log = log_df[log_df["routine_type"] == routine_type].copy()

    sequences = []

    for case_id in routine_log["case:concept:name"].unique():

        trace = routine_log[
            routine_log["case:concept:name"] == case_id
        ].sort_values("time:timestamp")

        # Use full activity labels directly
        seq = tuple(trace["concept:name"])
        sequences.append(seq)

    if not sequences:
        return []

    counts = Counter(sequences)
    canonical = counts.most_common(1)[0][0]

    print(f"\nRoutine Type: {routine_type}")
    print(f"  Total traces: {len(sequences)}")
    print(f"  Canonical sequence: {list(canonical)}")
    print(f"  Unique sequences found: {len(counts)}")

    return list(canonical)


# ============================================================
# Helper 2: Strict Trace Normalization (Duplicate-Safe)
# ============================================================

def normalize_trace(trace, canonical_sequence):
    """
    Strictly rebuilds a trace according to the canonical sequence.

    Guarantees:
    -----------
    - All traces follow identical ordering
    - Duplicate activities handled correctly
    - Full activity labels preserved (no stripping)
    - Trace duration preserved
    """

    trace = trace.copy().sort_values("time:timestamp").reset_index(drop=True)

    original_start = trace["time:timestamp"].min()
    original_end = trace["time:timestamp"].max()
    duration = (original_end - original_start).total_seconds()

    # --------------------------------------------------------
    # Build event pool: activity → queue of event rows
    # --------------------------------------------------------

    event_pool = defaultdict(deque)

    for _, row in trace.iterrows():
        activity = row["concept:name"]
        event_pool[activity].append(row)

    # --------------------------------------------------------
    # Strict rebuild following canonical order
    # --------------------------------------------------------

    normalized_rows = []

    for activity in canonical_sequence:

        if not event_pool[activity]:
            raise ValueError(
                f"Trace {trace['case:concept:name'].iloc[0]} "
                f"is missing required activity '{activity}'."
            )

        normalized_rows.append(event_pool[activity].popleft())

    normalized_trace = pd.DataFrame(normalized_rows)

    # --------------------------------------------------------
    # Timestamp reassignment preserving original duration
    # --------------------------------------------------------

    n = len(normalized_trace)

    if n == 1:
        normalized_trace["time:timestamp"] = [original_start]

    else:
        step = max(duration / (n - 1), 0.1)

        normalized_trace["time:timestamp"] = [
            original_start + pd.Timedelta(seconds=i * step)
            for i in range(n)
        ]

    return normalized_trace.reset_index(drop=True)


# ============================================================
# Main Function: Normalize Activity Sequences
# ============================================================

def normalize_activity_sequences(log_input, output_path=None, min_gap_seconds=1.0):
    print("I am heree....")
    """
    Normalize activity sequences so that all traces belonging to the same
    routine_type follow the exact same canonical event order.

    Features:
    ---------
    1. Extracts canonical sequence per routine type
    2. Strictly enforces ordering across traces (duplicate-safe)
    3. Preserves trace duration
    4. Prevents overlap between consecutive traces
    """

    # --------------------------------------------------------
    # Step 0: Load Log
    # --------------------------------------------------------

    if isinstance(log_input, str):
        log = pd.read_csv(log_input)

    elif isinstance(log_input, pd.DataFrame):
        log = log_input.copy()

    else:
        raise ValueError("log_input must be a file path or a DataFrame")

    required_columns = [
        "time:timestamp",
        "case:concept:name",
        "routine_type",
        "concept:name",
    ]

    missing = [c for c in required_columns if c not in log.columns]
    if missing:
        raise ValueError(f"Missing required columns: {missing}")

    log["time:timestamp"] = pd.to_datetime(log["time:timestamp"])

    print("=" * 70)
    print("NORMALIZING ACTIVITY SEQUENCES (STRICT CANONICAL ORDER)")
    print("=" * 70)

    # --------------------------------------------------------
    # Step 1: Canonical Sequence Per Routine Type
    # --------------------------------------------------------

    routine_types = sorted(log["routine_type"].unique())

    canonical_sequences = {
        rt: find_canonical_sequence(log, rt)
        for rt in routine_types
    }

    # --------------------------------------------------------
    # Step 2: Normalize Traces Sequentially (No Overlap)
    # --------------------------------------------------------

    normalized_traces = []

    for routine_type in routine_types:

        routine_log = log[log["routine_type"] == routine_type].copy()
        canonical_sequence = canonical_sequences[routine_type]

        if not canonical_sequence:
            print(f"Warning: No canonical sequence found for {routine_type}")
            normalized_traces.append(routine_log)
            continue

        # Sort traces by start time
        case_starts = (
            routine_log.groupby("case:concept:name")["time:timestamp"]
            .min()
            .sort_values()
        )

        print(f"\nNormalizing {len(case_starts)} traces for routine {routine_type}...")

        current_max_time = None
        normalized_routine = []

        for case_id in case_starts.index:

            trace = routine_log[routine_log["case:concept:name"] == case_id]

            # Normalize trace ordering strictly
            normalized_trace_df = normalize_trace(trace, canonical_sequence)

            trace_start = normalized_trace_df["time:timestamp"].min()
            trace_end = normalized_trace_df["time:timestamp"].max()

            # Prevent overlap
            if current_max_time is not None:

                required_start = current_max_time + pd.Timedelta(
                    seconds=min_gap_seconds
                )

                if trace_start < required_start:
                    shift = required_start - trace_start
                    normalized_trace_df["time:timestamp"] += shift
                    trace_end = normalized_trace_df["time:timestamp"].max()

            current_max_time = trace_end
            normalized_routine.append(normalized_trace_df)

        normalized_traces.append(
            pd.concat(normalized_routine, ignore_index=True)
        )

    # --------------------------------------------------------
    # Step 3: Merge All Routine Types
    # --------------------------------------------------------

    normalized_log = pd.concat(normalized_traces, ignore_index=True)
    normalized_log = normalized_log.sort_values("time:timestamp").reset_index(drop=True)

    print("\n" + "=" * 70)
    print("NORMALIZATION COMPLETE")
    print("=" * 70)

    # --------------------------------------------------------
    # Step 4: Save Output (Optional)
    # --------------------------------------------------------

    if output_path:
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        normalized_log.to_csv(output_path, index=False)
        print(f"✓ Normalized log saved to: {output_path}")

    return normalized_log




